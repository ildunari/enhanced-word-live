"""
Enhanced Search and Replace with Formatting
A more LLM-friendly approach that combines search/replace with formatting options
"""

async def enhanced_search_and_replace(filename: str, find_text: str, replace_text: str,
                                    apply_formatting: bool = False,
                                    bold: Optional[bool] = None, 
                                    italic: Optional[bool] = None,
                                    underline: Optional[bool] = None, 
                                    color: Optional[str] = None,
                                    font_size: Optional[int] = None, 
                                    font_name: Optional[str] = None,
                                    match_case: bool = True,
                                    whole_words_only: bool = False) -> str:
    """Enhanced search and replace with formatting options.
    
    Args:
        filename: Path to the Word document
        find_text: Text to search for
        replace_text: Text to replace with
        apply_formatting: Whether to apply formatting to the replaced text
        bold: Set replaced text bold (True/False)
        italic: Set replaced text italic (True/False)
        underline: Set replaced text underlined (True/False)
        color: Text color for replaced text (e.g., 'red', 'blue', etc.)
        font_size: Font size in points for replaced text
        font_name: Font name/family for replaced text
        match_case: Whether to match case (default True)
        whole_words_only: Whether to match whole words only (default False)
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Prepare search text based on options
        search_text = find_text if match_case else find_text.lower()
        
        count = 0
        
        # Enhanced find and replace with formatting
        count += _enhanced_replace_in_paragraphs(doc.paragraphs, find_text, replace_text, 
                                                apply_formatting, bold, italic, underline, 
                                                color, font_size, font_name, match_case, whole_words_only)
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    count += _enhanced_replace_in_paragraphs(cell.paragraphs, find_text, replace_text,
                                                           apply_formatting, bold, italic, underline,
                                                           color, font_size, font_name, match_case, whole_words_only)
        
        if count > 0:
            doc.save(filename)
            formatting_applied = " with formatting" if apply_formatting else ""
            return f"Replaced {count} occurrence(s) of '{find_text}' with '{replace_text}'{formatting_applied}."
        else:
            return f"No occurrences of '{find_text}' found."
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"


def _enhanced_replace_in_paragraphs(paragraphs, find_text, replace_text, apply_formatting,
                                   bold, italic, underline, color, font_size, font_name,
                                   match_case, whole_words_only):
    """Helper function to replace text in paragraphs with optional formatting."""
    import re
    from docx.shared import Pt, RGBColor
    
    count = 0
    
    for para in paragraphs:
        # Check if paragraph contains the search text
        para_text = para.text if match_case else para.text.lower()
        search_text = find_text if match_case else find_text.lower()
        
        if search_text in para_text:
            # Create regex pattern for replacement
            if whole_words_only:
                pattern = r'\b' + re.escape(find_text) + r'\b'
            else:
                pattern = re.escape(find_text)
            
            flags = 0 if match_case else re.IGNORECASE
            
            # Process each run in the paragraph
            for run in para.runs:
                if re.search(pattern, run.text, flags):
                    # Count matches before replacement
                    matches = len(re.findall(pattern, run.text, flags))
                    count += matches
                    
                    # Replace the text
                    run.text = re.sub(pattern, replace_text, run.text, flags=flags)
                    
                    # Apply formatting if requested
                    if apply_formatting:
                        if bold is not None:
                            run.bold = bold
                        if italic is not None:
                            run.italic = italic
                        if underline is not None:
                            run.underline = underline
                        if color:
                            _apply_color_to_run(run, color)
                        if font_size:
                            run.font.size = Pt(font_size)
                        if font_name:
                            run.font.name = font_name
    
    return count


def _apply_color_to_run(run, color):
    """Apply color to a run with error handling."""
    from docx.shared import RGBColor
    
    # Define common RGB colors
    color_map = {
        'red': RGBColor(255, 0, 0),
        'blue': RGBColor(0, 0, 255),
        'green': RGBColor(0, 128, 0),
        'yellow': RGBColor(255, 255, 0),
        'black': RGBColor(0, 0, 0),
        'gray': RGBColor(128, 128, 128),
        'white': RGBColor(255, 255, 255),
        'purple': RGBColor(128, 0, 128),
        'orange': RGBColor(255, 165, 0),
        'brown': RGBColor(165, 42, 42),
        'pink': RGBColor(255, 192, 203),
        'cyan': RGBColor(0, 255, 255),
        'magenta': RGBColor(255, 0, 255),
        'lime': RGBColor(0, 255, 0),
        'navy': RGBColor(0, 0, 128),
        'maroon': RGBColor(128, 0, 0),
        'olive': RGBColor(128, 128, 0),
        'teal': RGBColor(0, 128, 128)
    }
    
    try:
        if color.lower() in color_map:
            run.font.color.rgb = color_map[color.lower()]
        else:
            # Try to parse as hex color (e.g., "#FF0000")
            if color.startswith('#') and len(color) == 7:
                hex_color = color[1:]
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            else:
                # Default to black if color not recognized
                run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        # Fallback to black
        run.font.color.rgb = RGBColor(0, 0, 0)


async def format_specific_words(filename: str, word_list: List[str], 
                               bold: Optional[bool] = None,
                               italic: Optional[bool] = None,
                               underline: Optional[bool] = None,
                               color: Optional[str] = None,
                               font_size: Optional[int] = None,
                               font_name: Optional[str] = None,
                               match_case: bool = True,
                               whole_words_only: bool = True) -> str:
    """Format specific words throughout the document.
    
    Args:
        filename: Path to the Word document
        word_list: List of words to format
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (e.g., 'red', 'blue', etc.)
        font_size: Font size in points
        font_name: Font name/family
        match_case: Whether to match case (default True)
        whole_words_only: Whether to match whole words only (default True)
    """
    total_count = 0
    results = []
    
    for word in word_list:
        # Use enhanced search and replace with same text for find and replace
        result = await enhanced_search_and_replace(
            filename=filename,
            find_text=word,
            replace_text=word,  # Same text, just apply formatting
            apply_formatting=True,
            bold=bold,
            italic=italic,
            underline=underline,
            color=color,
            font_size=font_size,
            font_name=font_name,
            match_case=match_case,
            whole_words_only=whole_words_only
        )
        results.append(f"'{word}': {result}")
    
    return "\n".join(results)


# Example usage for your PCL research paper:

async def format_research_paper_terms(filename: str) -> str:
    """Format common research terms in a PCL paper with appropriate styling."""
    
    # Format drug names in blue and bold
    drug_names = ["dolutegravir", "meloxicam", "dexamethasone", "DTG", "MLX", "DEX"]
    await format_specific_words(filename, drug_names, bold=True, color="blue")
    
    # Format polymer terms in green
    polymer_terms = ["polycaprolactone", "PCL", "mesophase", "crystallinity"]
    await format_specific_words(filename, polymer_terms, color="green")
    
    # Format statistical terms in red and italic
    stats_terms = ["p < 0.05", "significant", "correlation", "ANOVA"]
    await format_specific_words(filename, stats_terms, italic=True, color="red")
    
    # Format temperature values in orange
    temp_terms = ["25°C", "50°C"]
    await format_specific_words(filename, temp_terms, color="orange")
    
    return "Research paper terms formatted successfully!"