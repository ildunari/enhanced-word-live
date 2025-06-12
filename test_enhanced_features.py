#!/usr/bin/env python3
"""
Test script for enhanced Word MCP server features.
Tests the new Tier 1 features: Enhanced Search/Replace, Review Tools, and Section Management.
"""

import asyncio
import os
from pathlib import Path
from docx import Document

# Import our enhanced tools
from word_document_server.tools.content_tools import (
    enhanced_search_and_replace, 
    format_specific_words,
    format_research_paper_terms
)
from word_document_server.tools.review_tools import (
    extract_comments,
    extract_track_changes,
    generate_review_summary,
    add_comment
)
from word_document_server.tools.section_tools import (
    extract_sections_by_heading,
    extract_section_content,
    generate_table_of_contents,
    get_section_statistics
)


async def create_test_document():
    """Create a test document for demonstrating enhanced features."""
    filename = "test_enhanced_features.docx"
    
    # Create a document with academic content
    doc = Document()
    
    # Add title
    title = doc.add_heading('PCL Mesophase Drug Delivery Research', 0)
    
    # Add abstract section
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        'This study investigates the use of polycaprolactone (PCL) mesophases for controlled '
        'drug delivery of three compounds: dolutegravir (DTG), meloxicam (MLX), and '
        'dexamethasone (DEX). Statistical analysis revealed significant correlations '
        '(p < 0.05) between mesophase content and release kinetics at both 25°C and 50°C.'
    )
    
    # Add introduction section
    doc.add_heading('Introduction', level=1)
    doc.add_paragraph(
        'Polycaprolactone has emerged as a promising biodegradable polymer for pharmaceutical '
        'applications. The crystallinity of PCL can be modulated through thermomechanical '
        'processing to create mesophase structures that influence drug release profiles.'
    )
    
    doc.add_heading('Drug Compounds', level=2)
    doc.add_paragraph(
        'Three model drugs were selected: dolutegravir for HIV treatment, '
        'meloxicam as an anti-inflammatory agent, and dexamethasone as a corticosteroid. '
        'Each compound exhibits different solubility characteristics affecting release kinetics.'
    )
    
    # Add methods section
    doc.add_heading('Methods', level=1)
    doc.add_paragraph(
        'Samples were processed using compression molding at temperatures of 25°C and 50°C. '
        'X-ray diffraction analysis was performed to quantify crystallinity changes. '
        'Release studies were conducted in phosphate-buffered saline with ANOVA statistical analysis.'
    )
    
    # Add results section
    doc.add_heading('Results', level=1)
    doc.add_paragraph(
        'Significant differences were observed between treatment groups. '
        'The correlation between mesophase content and drug release was highly significant '
        'with r² values exceeding 0.85 for all compounds tested.'
    )
    
    # Add table
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # Add table headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Compound'
    hdr_cells[1].text = 'Mesophase %'
    hdr_cells[2].text = 'Release Rate (μg/mL/h)'
    
    # Add data
    data = [
        ['DTG', '23.5 ± 2.1', '15.2 ± 1.8'],
        ['MLX', '31.7 ± 3.4', '22.1 ± 2.5'],
        ['DEX', '18.9 ± 1.9', '12.7 ± 1.4']
    ]
    
    for i, row_data in enumerate(data, 1):
        row_cells = table.rows[i].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
    
    # Add conclusion
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'This research demonstrates the potential of PCL mesophase engineering for '
        'precise control of drug release kinetics. The significant correlations observed '
        'support the hypothesis that processing-induced mesophases can be leveraged '
        'for pharmaceutical applications.'
    )
    
    doc.save(filename)
    return filename


async def test_enhanced_search_replace():
    """Test the enhanced search and replace functionality."""
    print("\n=== TESTING ENHANCED SEARCH AND REPLACE ===")
    
    filename = await create_test_document()
    
    # Test 1: Basic enhanced search and replace with formatting
    print("Test 1: Replace 'PCL' with formatted version...")
    result = await enhanced_search_and_replace(
        filename=filename,
        find_text="PCL",
        replace_text="PCL",
        apply_formatting=True,
        bold=True,
        color="blue"
    )
    print(f"Result: {result}")
    
    # Test 2: Format specific research terms
    print("\nTest 2: Format research paper terms...")
    result = await format_research_paper_terms(filename)
    print(f"Result: {result}")
    
    # Test 3: Format statistical terms
    print("\nTest 3: Format statistical significance values...")
    result = await format_specific_words(
        filename=filename,
        word_list=["p < 0.05", "r²", "±"],
        bold=True,
        color="red",
        whole_words_only=False
    )
    print(f"Result: {result}")
    
    return filename


async def test_review_tools(filename):
    """Test the review and collaboration tools."""
    print("\n=== TESTING REVIEW TOOLS ===")
    
    # Test 1: Add a comment
    print("Test 1: Adding a comment to the abstract...")
    result = await add_comment(
        filename=filename,
        paragraph_index=2,  # Abstract paragraph
        comment_text="Consider adding more details about the mechanism of mesophase formation.",
        author="Dr. Smith"
    )
    print(f"Result: {result}")
    
    # Test 2: Extract comments (if any exist)
    print("\nTest 2: Extracting comments...")
    result = await extract_comments(filename)
    print(f"Result: {result}")
    
    # Test 3: Extract track changes (if any exist)
    print("\nTest 3: Extracting track changes...")
    result = await extract_track_changes(filename)
    print(f"Result: {result}")
    
    # Test 4: Generate review summary
    print("\nTest 4: Generating review summary...")
    result = await generate_review_summary(filename)
    print(f"Result: {result}")


async def test_section_tools(filename):
    """Test the section management tools."""
    print("\n=== TESTING SECTION MANAGEMENT TOOLS ===")
    
    # Test 1: Extract sections by heading
    print("Test 1: Extracting document sections...")
    result = await extract_sections_by_heading(filename)
    print(f"Result: {result}")
    
    # Test 2: Extract specific section content
    print("\nTest 2: Extracting 'Methods' section content...")
    result = await extract_section_content(filename, "Methods")
    print(f"Result: {result}")
    
    # Test 3: Generate table of contents
    print("\nTest 3: Generating table of contents...")
    result = await generate_table_of_contents(filename)
    print(f"Result: {result}")
    
    # Test 4: Get section statistics
    print("\nTest 4: Getting section statistics...")
    result = await get_section_statistics(filename)
    print(f"Result: {result}")


async def run_all_tests():
    """Run all enhanced feature tests."""
    print("=" * 60)
    print("TESTING ENHANCED WORD MCP SERVER FEATURES")
    print("=" * 60)
    
    try:
        # Test enhanced search and replace
        filename = await test_enhanced_search_replace()
        
        # Test review tools
        await test_review_tools(filename)
        
        # Test section tools
        await test_section_tools(filename)
        
        print("\n" + "=" * 60)
        print("ALL TESTS COMPLETED SUCCESSFULLY!")
        print(f"Test document created: {filename}")
        print("=" * 60)
        
    except Exception as e:
        print(f"\nERROR DURING TESTING: {str(e)}")
        import traceback
        traceback.print_exc()


if __name__ == "__main__":
    asyncio.run(run_all_tests())