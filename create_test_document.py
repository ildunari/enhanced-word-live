#!/usr/bin/env python3
"""
Create a test Word document to verify regex batch processing issues.
This script creates a document with multiple curly brace instances that
should reveal the text corruption issues in the enhanced_search_and_replace function.
"""

from docx import Document
from docx.shared import Inches

def create_test_document():
    """Create a test document with various curly brace scenarios."""
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Regex Batch Processing Test Document', 0)
    
    # Test Case 1: Simple paragraph with multiple curly braces
    doc.add_heading('Test Case 1: Multiple instances in single paragraph', level=1)
    p1 = doc.add_paragraph()
    p1.add_run("Please review ").bold = False
    p1.add_run("{section 2.1}").bold = False
    p1.add_run(" and ").bold = False  
    p1.add_run("{appendix A}").bold = False
    p1.add_run(" carefully.").bold = False
    
    # Test Case 2: Multiple paragraphs with curly braces
    doc.add_heading('Test Case 2: Multiple paragraphs', level=1)
    doc.add_paragraph("First paragraph contains {important data} for analysis.")
    doc.add_paragraph("Second paragraph has {critical information} and {key findings}.")
    doc.add_paragraph("Third paragraph mentions {final results} at the end.")
    
    # Test Case 3: Complex formatting with curly braces
    doc.add_heading('Test Case 3: Complex formatting', level=1)
    p3 = doc.add_paragraph()
    p3.add_run("The document discusses ").bold = False
    run1 = p3.add_run("{research methodology}")
    run1.bold = True  # Pre-existing bold formatting
    p3.add_run(" and also covers ").bold = False
    run2 = p3.add_run("{data analysis}")
    run2.italic = True  # Pre-existing italic formatting
    p3.add_run(" in detail.").bold = False
    
    # Test Case 4: Edge cases
    doc.add_heading('Test Case 4: Edge cases', level=1)
    doc.add_paragraph("Empty braces {} should be handled.")
    doc.add_paragraph("Nested {braces {inside} braces} are tricky.")
    doc.add_paragraph("Multiple {a} {b} {c} {d} instances in sequence.")
    
    # Test Case 5: Long text to potentially span multiple runs
    doc.add_heading('Test Case 5: Long text scenarios', level=1)
    long_text = ("This is a very long paragraph that contains "
                "{very long text inside curly braces that might span multiple runs "
                "depending on how Word handles the document formatting and text flow} "
                "and continues with more text after the braces.")
    doc.add_paragraph(long_text)
    
    # Save the document
    filename = "regex_test_document.docx"
    doc.save(filename)
    print(f"✅ Test document created: {filename}")
    
    return filename

if __name__ == "__main__":
    create_test_document()