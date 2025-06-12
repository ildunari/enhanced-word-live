#!/usr/bin/env python3
"""
Test script to verify both the main project and copy folder have the fix applied.
This script will test both implementations to ensure consistency.
"""

import sys
import os
from docx import Document
from pathlib import Path
import shutil
import asyncio

def test_version(project_path, version_name):
    """Test a specific version of the enhanced_search_and_replace function."""
    print(f"\n{'='*60}")
    print(f"Testing {version_name}")
    print(f"Path: {project_path}")
    print('='*60)
    
    # Add the project directory to Python path
    sys.path.insert(0, str(project_path))
    
    try:
        # Import the function from this version
        from word_document_server.tools.content_tools import enhanced_search_and_replace
        
        # Create a test document copy for this version
        test_file = "regex_test_document.docx"
        if not os.path.exists(test_file):
            print(f"❌ Test file {test_file} not found.")
            return False
        
        test_copy = f"regex_test_{version_name.lower().replace(' ', '_')}.docx"
        shutil.copy2(test_file, test_copy)
        
        print(f"📄 Using test file: {test_copy}")
        
        # Read initial state
        doc_before = Document(test_copy)
        para_before = None
        for para in doc_before.paragraphs:
            if "Please review {section 2.1} and {appendix A} carefully." in para.text:
                para_before = para.text
                break
        
        print(f"   Before: '{para_before}'")
        
        # Run the test
        async def run_test():
            return await enhanced_search_and_replace(
                filename=test_copy,
                find_text=r"\{[^}]+\}",
                replace_text=r"\g<0>",
                use_regex=True,
                apply_formatting=True,
                bold=True,
                match_case=True
            )
        
        result = asyncio.run(run_test())
        print(f"   Result: {result}")
        
        # Read final state
        doc_after = Document(test_copy)
        para_after = None
        for para in doc_after.paragraphs:
            if "Please review" in para.text and "{section 2.1}" in para.text:
                para_after = para.text
                break
        
        print(f"   After:  '{para_after}'")
        
        # Check if it's fixed (text structure preserved)
        if para_before == para_after:
            print(f"✅ {version_name}: CORRUPTION FIXED!")
            return True
        else:
            print(f"❌ {version_name}: Still corrupted")
            return False
            
    except Exception as e:
        print(f"❌ {version_name}: Error during test: {e}")
        import traceback
        traceback.print_exc()
        return False
    finally:
        # Remove this path from sys.path to avoid conflicts
        if str(project_path) in sys.path:
            sys.path.remove(str(project_path))

def main():
    """Test both versions."""
    print("🔍 Testing both implementations for regex corruption fixes...")
    
    main_project = Path("/Users/kosta/Documents/ProjectsCode/kosta-enhanced-word-mcp-server")
    copy_project = Path("/Users/kosta/Documents/ProjectsCode/kosta-enhanced-word-mcp-server copy")
    
    # Test main project
    main_success = test_version(main_project, "Main Project")
    
    # Test copy project
    copy_success = test_version(copy_project, "Copy Project")
    
    # Summary
    print(f"\n{'='*60}")
    print("🎯 FINAL RESULTS")
    print('='*60)
    print(f"Main Project: {'✅ FIXED' if main_success else '❌ BROKEN'}")
    print(f"Copy Project: {'✅ FIXED' if copy_success else '❌ BROKEN'}")
    
    if main_success and copy_success:
        print("\n🎉 SUCCESS: Both implementations are fixed! ✅")
        return True
    else:
        print("\n⚠️  Some implementations still have issues")
        return False

if __name__ == "__main__":
    success = main()
    sys.exit(0 if success else 1)