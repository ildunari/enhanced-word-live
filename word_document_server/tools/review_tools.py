"""
Review tools for Word Document Server - Tier 1 Feature.

These tools handle collaboration features including comments, track changes,
and review management for academic research workflows.
"""
import os
from typing import List, Optional, Dict, Any
from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor
from xml.etree import ElementTree as ET

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.session_utils import resolve_document_path

class WordDocumentError(Exception):
    """Base exception for Word document operations."""
    pass

class DocumentNotFoundError(WordDocumentError):
    """Raised when a document file is not found."""
    pass

class DocumentAccessError(WordDocumentError):
    """Raised when a document cannot be accessed or is locked."""
    pass

class DocumentCorruptionError(WordDocumentError):
    """Raised when a document appears to be corrupted."""
    pass

class InvalidPathError(WordDocumentError):
    """Raised when an invalid file path is provided."""
    pass

def manage_comments(
    document_id: str = None,
    filename: str = None,
    action: str = "list",
    paragraph_index: int = None,
    comment_text: str = None,
    author: str = None,
    comment_id: str = None
) -> str:
    """Enhanced comment management with extraction, creation, and resolution capabilities.
    
    This enhanced function now provides complete comment lifecycle management including extraction,
    creation, resolution, and deletion while maintaining backward compatibility for listing comments.
    
    Args:
        document_id: Session document ID (preferred)
        filename: Path to the Word document (legacy, for backward compatibility)
        action: Operation to perform:
            - "list" (default): Extract all comments with metadata
            - "add": Create new comment on specified paragraph  
            - "resolve": Mark comment as resolved
            - "delete": Remove comment completely
        paragraph_index: Zero-based paragraph index for new comments (required for "add")
        comment_text: Comment content text (required for "add")
        author: Comment author name (optional, defaults to "User")
        comment_id: Identifier for existing comment operations (required for "resolve"/"delete")
    
    Returns:
        Formatted string with comment information or operation status
    """
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Resolve document path from document_id or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    # Validate action parameter
    valid_actions = ["list", "add", "resolve", "delete"]
    if action not in valid_actions:
        return f"Invalid action: {action}. Must be one of: {', '.join(valid_actions)}"
    
    # Validate required parameters for each action
    if action == "add":
        if paragraph_index is None:
            return "Parameter 'paragraph_index' is required for action 'add'"
        if not comment_text:
            return "Parameter 'comment_text' is required for action 'add'"
    
    if action in ["resolve", "delete"]:
        if not comment_id:
            return f"Parameter 'comment_id' is required for action '{action}'"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Handle comment management actions
    if action in ["add", "resolve", "delete"]:
        # Check if file is writeable for modification actions
        is_writeable, error_message = check_file_writeable(filename)
        if not is_writeable:
            return f"Cannot modify document: {error_message}"
    
    try:
        doc = Document(filename)
        
        # Handle comment management actions
        if action == "add":
            # Add new comment to specified paragraph
            if paragraph_index >= len(doc.paragraphs):
                return f"Paragraph index {paragraph_index} is out of range (document has {len(doc.paragraphs)} paragraphs)"
            
            paragraph = doc.paragraphs[paragraph_index]
            author_name = author or "User"
            
            # Add comment as a text annotation (simplified implementation)
            import uuid
            comment_uuid = str(uuid.uuid4())[:8]
            comment_marker = f" [COMMENT-{comment_uuid} by {author_name}: {comment_text}]"
            
            # Add the comment text to the end of the paragraph
            if paragraph.text:
                paragraph.text += comment_marker
            else:
                paragraph.text = comment_marker
            
            doc.save(filename)
            return f"Successfully added comment {comment_uuid} to paragraph {paragraph_index}"
        
        elif action in ["resolve", "delete"]:
            # Search through document for comment markers
            comment_found = False
            
            for para_idx, paragraph in enumerate(doc.paragraphs):
                if f"[COMMENT-{comment_id}" in paragraph.text or f"[RESOLVED-{comment_id}" in paragraph.text:
                    comment_found = True
                    
                    if action == "resolve":
                        # Mark as resolved
                        paragraph.text = paragraph.text.replace(f"[COMMENT-{comment_id}", f"[RESOLVED-{comment_id}")
                        doc.save(filename)
                        return f"Successfully resolved comment {comment_id}"
                    
                    elif action == "delete":
                        # Remove comment completely
                        start_markers = [f"[COMMENT-{comment_id}", f"[RESOLVED-{comment_id}"]
                        for start_marker in start_markers:
                            if start_marker in paragraph.text:
                                start_pos = paragraph.text.find(start_marker)
                                if start_pos != -1:
                                    end_pos = paragraph.text.find("]", start_pos)
                                    if end_pos != -1:
                                        comment_part = paragraph.text[start_pos:end_pos+1]
                                        paragraph.text = paragraph.text.replace(comment_part, "")
                                        doc.save(filename)
                                        return f"Successfully deleted comment {comment_id}"
                    break
            
            if not comment_found:
                return f"Comment {comment_id} not found in document"
        
        # Handle list action - search for text-based comment markers
        elif action == "list":
            comments_info = []
            
            # Search through all paragraphs for comment markers
            for para_idx, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text
                
                # Find all comment markers in this paragraph
                import re
                # Pattern matches: [COMMENT-12345678 by Author: comment text] or [RESOLVED-12345678 by Author: comment text]
                pattern = r'\[(COMMENT|RESOLVED)-([a-f0-9]{8}) by ([^:]+): ([^\]]+)\]'
                matches = re.findall(pattern, text)
                
                for match in matches:
                    status, comment_id, author_name, comment_content = match
                    comments_info.append({
                        'id': comment_id,
                        'author': author_name,
                        'status': status.lower(),  # 'comment' or 'resolved'
                        'text': comment_content,
                        'paragraph_index': para_idx
                    })
            
            if not comments_info:
                return "No comments found in the document."
            
            # Format output
            result = f"Found {len(comments_info)} comments:\n\n"
            for i, comment in enumerate(comments_info, 1):
                status_indicator = " (RESOLVED)" if comment['status'] == 'resolved' else ""
                result += f"Comment {i} (ID: {comment['id']}){status_indicator}:\n"
                result += f"  Author: {comment['author']}\n"
                result += f"  Paragraph: {comment['paragraph_index']}\n"
                result += f"  Text: {comment['text']}\n\n"
            
            return result
    
    except Exception as e:
        return f"Failed to manage comments: {str(e)}"


def extract_track_changes(document_id: str = None, filename: str = None) -> str:
    """Extract track changes information from a Word document.
    
    Args:
        document_id (str, optional): Session document identifier (preferred)
        filename (str, optional): Path to the Word document
    
    Returns:
        Formatted string with all track changes, authors, and change types
    """
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Resolve document path from session or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        doc = Document(filename)
        changes_info = []
        
        # Access the document's XML to extract revision information
        document_xml = doc.element.xml
        root = ET.fromstring(document_xml)
        
        # Extract track changes with namespace handling
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        # Find insertions
        for ins in root.findall('.//w:ins', ns):
            author = ins.get(qn('w:author'), 'Unknown')
            date = ins.get(qn('w:date'), 'Unknown')
            change_id = ins.get(qn('w:id'), 'Unknown')
            
            # Extract inserted text
            inserted_text = ""
            for text_elem in ins.findall('.//w:t', ns):
                if text_elem.text:
                    inserted_text += text_elem.text
            
            changes_info.append({
                'type': 'insertion',
                'id': change_id,
                'author': author,
                'date': date,
                'text': inserted_text
            })
        
        # Find deletions
        for del_elem in root.findall('.//w:del', ns):
            author = del_elem.get(qn('w:author'), 'Unknown')
            date = del_elem.get(qn('w:date'), 'Unknown')
            change_id = del_elem.get(qn('w:id'), 'Unknown')
            
            # Extract deleted text
            deleted_text = ""
            for text_elem in del_elem.findall('.//w:delText', ns):
                if text_elem.text:
                    deleted_text += text_elem.text
            
            changes_info.append({
                'type': 'deletion',
                'id': change_id,
                'author': author,
                'date': date,
                'text': deleted_text
            })
        
        if not changes_info:
            return "No track changes found in the document."
        
        # Format output
        result = f"Found {len(changes_info)} track changes:\n\n"
        for i, change in enumerate(changes_info, 1):
            result += f"Change {i} (ID: {change['id']}):\n"
            result += f"  Type: {change['type'].title()}\n"
            result += f"  Author: {change['author']}\n"
            result += f"  Date: {change['date']}\n"
            result += f"  Text: '{change['text']}'\n\n"
        
        return result
    
    except Exception as e:
        return f"Failed to extract track changes: {str(e)}"


async def generate_review_summary(document_id: str = None, filename: str = None) -> str:
    """Generate a comprehensive review summary including comments and track changes.
    
    Args:
        document_id (str, optional): Session document identifier (preferred)
        filename (str, optional): Path to the Word document
    
    Returns:
        Formatted summary of all review elements suitable for academic collaboration
    """
    # Resolve document path from session or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        # Get comments
        comments_result = manage_comments(filename, action="list")
        
        # Get track changes
        changes_result = extract_track_changes(filename)
        
        # Generate summary
        summary = f"=== REVIEW SUMMARY FOR {os.path.basename(filename)} ===\n\n"
        summary += "COMMENTS:\n"
        summary += "-" * 50 + "\n"
        summary += comments_result + "\n\n"
        
        summary += "TRACK CHANGES:\n"
        summary += "-" * 50 + "\n"
        summary += changes_result + "\n\n"
        
        summary += "=== END REVIEW SUMMARY ===\n"
        
        return summary
    
    except Exception as e:
        return f"Failed to generate review summary: {str(e)}"


async def manage_track_changes(
    document_id: str = None,
    filename: str = None,
    action: str = None,
    change_ids: Optional[List[str]] = None,
    author_filter: Optional[str] = None
) -> str:
    """Unified track changes management function for comprehensive revision control.
    
    This function consolidates all track changes operations into a single comprehensive tool.
    It replaces accept_all_changes and reject_all_changes with enhanced selective capabilities
    for granular change management in collaborative document workflows.
    
    Args:
        document_id (str): Session document ID (preferred)
        filename (str): Path to the Word document (legacy, for backward compatibility)
        
        action (str): Track changes action to perform:
            - "accept_all": Accept all tracked changes in document
            - "reject_all": Reject all tracked changes in document  
            - "accept_selective": Accept only specific changes (requires change_ids or author_filter)
            - "reject_selective": Reject only specific changes (requires change_ids or author_filter)
        
        change_ids (List[str], optional): Specific change identifiers for selective operations
            - Used with action="accept_selective" or "reject_selective"
            - Each ID corresponds to a specific tracked change
            - Can be obtained from document analysis tools
            - Example: ["change_1", "change_5", "change_12"]
        
        author_filter (str, optional): Process changes only from specific author
            - Case-sensitive author name matching
            - Works with both bulk and selective operations
            - Useful for reviewing contributions from specific collaborators
            - Example: "Dr. Jane Smith" or "john.doe@company.com"
    
    Returns:
        str: Status message describing operation result:
            - Success: "Successfully {action} {count} changes"
            - Error: Specific error message with troubleshooting guidance
            - Warning: Information about partially completed operations
    
    Use Cases:
        📝 Final Document Preparation: Accept all changes before publication
        👥 Collaborative Review: Selectively accept/reject reviewer suggestions
        🔄 Version Control: Manage changes from multiple authors systematically
        📋 Quality Control: Review and approve changes by expertise area
        ✅ Editorial Workflow: Process editorial changes in controlled manner
        🚫 Change Rollback: Reject unwanted or incorrect modifications
    
    Examples:
        # Accept all tracked changes for final document
        result = await manage_track_changes(document_id="final_report", action="accept_all")
        # Returns: "Successfully accepted 47 changes"
        
        # Reject all changes to restore original version
        result = await manage_track_changes(document_id="draft", action="reject_all")
        # Returns: "Successfully rejected 23 changes"
        
        # Accept only changes from lead researcher
        result = await manage_track_changes(document_id="research_paper", action="accept_selective", 
                                           author_filter="Dr. Sarah Johnson")
        # Returns: "Successfully accepted 12 changes by Dr. Sarah Johnson"
        
        # Reject changes from specific reviewer
        result = await manage_track_changes(document_id="manuscript", action="reject_selective",
                                           author_filter="External Reviewer")
        # Returns: "Successfully rejected 8 changes by External Reviewer"
        
        # Accept specific changes by ID (advanced usage)
        result = await manage_track_changes(document_id="document", action="accept_selective",
                                           change_ids=["change_5", "change_18", "change_23"])
        # Returns: "Successfully accepted 3 specific changes"
        
        # Process all changes from multiple authors
        result = await manage_track_changes(document_id="collaborative_doc", action="accept_all",
                                           author_filter="Editor Team")
        # Returns: "Successfully accepted 15 changes by Editor Team"
        
        # Reject all editorial suggestions
        result = await manage_track_changes(document_id="author_manuscript", action="reject_all",
                                           author_filter="Copy Editor")
        # Returns: "Successfully rejected 31 changes by Copy Editor"
    
    Error Handling:
        - Document not found: "Document '{document_id}' not found in sessions"
        - File not writable: "Cannot modify document: {reason}. Consider creating a copy first."
        - Invalid action: "Invalid action: {action}. Must be one of: accept_all, reject_all, accept_selective, reject_selective"
        - Missing parameters: "Invalid parameter: either change_ids or author_filter required for selective operations"
        - No changes found: "No tracked changes found in document"
        - Author not found: "No changes found by author: {author_filter}"
        - Document corruption: "Error processing changes: {error_details}"
        - Protection conflict: "Document is protected and changes cannot be processed"
    
    Workflow Integration:
        1. Document Analysis: Use get_text with search to identify areas needing review
        2. Author Review: Filter changes by author_filter to review specific contributions
        3. Selective Processing: Use change_ids for granular control over specific edits
        4. Bulk Operations: Use accept_all/reject_all for final document preparation
        5. Version Control: Combine with document protection for controlled workflows
    
    Performance Notes:
        - Large documents with many changes may take longer to process
        - Selective operations are generally faster than bulk operations
        - Author filtering is more efficient than change ID filtering
        - Consider processing in batches for documents with hundreds of changes
    
    Security Considerations:
        - Requires write access to document file
        - Changes are permanently applied and cannot be undone without backup
        - Document protection must be removed before processing changes
        - Maintains document integrity and formatting during change processing
    """
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Resolve document path from document_id or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    # Validate required parameters
    if not action:
        return "Error: action parameter is required"
    
    # Validate action parameter
    valid_actions = ["accept_all", "reject_all", "accept_selective", "reject_selective"]
    if action not in valid_actions:
        return f"Invalid action: {action}. Must be one of: {', '.join(valid_actions)}"
    
    # Validate selective action parameters
    if action in ["accept_selective", "reject_selective"] and not change_ids and not author_filter:
        return "Invalid parameter: either change_ids or author_filter is required for selective operations"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        document_xml = doc.element
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        changes_processed = 0
        
        if action in ["accept_all", "accept_selective"]:
            # Accept changes logic - preserve the original XML manipulation
            
            # Process deletion markup (remove but preserve context)
            del_elements = document_xml.findall('.//w:del', ns)
            for del_elem in del_elements:
                # Check author filter if specified
                if author_filter and del_elem.get(qn('w:author')) != author_filter:
                    continue
                    
                del_elem.getparent().remove(del_elem)
                changes_processed += 1
            
            # Process insertion markup (keep inserted text, remove markup)
            ins_elements = document_xml.findall('.//w:ins', ns)
            for ins_elem in ins_elements:
                # Check author filter if specified
                if author_filter and ins_elem.get(qn('w:author')) != author_filter:
                    continue
                    
                # Move children out of ins element
                parent = ins_elem.getparent()
                for child in ins_elem:
                    parent.insert(list(parent).index(ins_elem), child)
                parent.remove(ins_elem)
                changes_processed += 1
        
        elif action in ["reject_all", "reject_selective"]:
            # Reject changes logic - preserve the original XML manipulation
            
            # Process insertion markup (remove inserted text)
            ins_elements = document_xml.findall('.//w:ins', ns)
            for ins_elem in ins_elements:
                # Check author filter if specified
                if author_filter and ins_elem.get(qn('w:author')) != author_filter:
                    continue
                    
                ins_elem.getparent().remove(ins_elem)
                changes_processed += 1
            
            # Process deletion markup (restore deleted text)
            del_elements = document_xml.findall('.//w:del', ns)
            for del_elem in del_elements:
                # Check author filter if specified
                if author_filter and del_elem.get(qn('w:author')) != author_filter:
                    continue
                    
                # Convert delText back to regular text
                for del_text in del_elem.findall('.//w:delText', ns):
                    # Create new text element
                    text_elem = ET.Element(qn('w:t'))
                    text_elem.text = del_text.text
                    
                    # Create new run
                    run_elem = ET.Element(qn('w:r'))
                    run_elem.append(text_elem)
                    
                    # Insert before deletion
                    parent = del_elem.getparent()
                    parent.insert(list(parent).index(del_elem), run_elem)
                
                # Remove the deletion element
                del_elem.getparent().remove(del_elem)
                changes_processed += 1
        
        doc.save(filename)
        
        # Build response message
        action_past_tense = {
            "accept_all": "accepted",
            "reject_all": "rejected", 
            "accept_selective": "accepted",
            "reject_selective": "rejected"
        }
        
        action_verb = action_past_tense[action]
        
        if author_filter:
            return f"All track changes by '{author_filter}' {action_verb} in {filename}. {changes_processed} changes processed."
        elif action in ["accept_all", "reject_all"]:
            return f"All track changes {action_verb} in {filename}. {changes_processed} changes processed."
        else:
            return f"Selected track changes {action_verb} in {filename}. {changes_processed} changes processed."
    
    except Exception as e:
        return f"Failed to {action.replace('_', ' ')}: {str(e)}"

