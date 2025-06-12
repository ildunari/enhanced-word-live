"""
Section management tools for Word Document Server - Tier 1 Feature.

These tools handle section organization via heading styles, perfect for
academic research document management and thesis synthesis.
"""
import os
import json
from typing import List, Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.session_utils import resolve_document_path


async def get_sections(
    document_id: str = None,
    filename: str = None,
    mode: str = "overview",
    section_title: Optional[str] = None,
    max_level: int = 3,
    include_subsections: bool = True,
    full_content: bool = False,
    case_sensitive: bool = False,
    output_format: str = "text",
    include_formatting: bool = False,
    formatting_detail: str = "basic"
) -> str:
    """Unified section extraction function for comprehensive document structure analysis.
    
    This function consolidates document structure analysis and content extraction into a single
    comprehensive tool. It replaces extract_sections_by_heading and extract_section_content
    with enhanced filtering, formatting, and output options for academic and professional workflows.
    
    Args:
        document_id (str, optional): Session document identifier (preferred)
        filename (str, optional): Path to the Word document (.docx extension added if missing)
            - Document should have heading-based structure for optimal results
            - Works with any document but most useful with proper heading hierarchy
        
        mode (str): Analysis and extraction mode:
            - "overview": Generate document structure outline (default)
            - "content": Extract full content from sections
            - Overview shows hierarchy, content provides detailed text
        
        section_title (str, optional): Target specific section by title
            - None: Process all sections (default)
            - Exact match: Find section with exact title
            - Partial match: Use case_sensitive parameter for control
            - Example: "Introduction", "3.2 Methodology", "Conclusion"
        
        max_level (int): Maximum heading level to include in results (1-9, default: 3)
            - 1: Only main chapters/sections
            - 3: Include subsections and sub-subsections
            - 9: Include all heading levels
            - Higher levels may produce very detailed output
        
        include_subsections (bool): Whether to include content from subsections (default: True)
            - True: Include all nested subsections within target section
            - False: Only include direct content of target section
            - Affects content mode primarily
        
        full_content (bool): Content detail level for overview mode (default: False)
            - False: Show structure with brief content summaries
            - True: Include full paragraph content in overview
            - Only affects overview mode, content mode always shows full content
        
        case_sensitive (bool): Case sensitivity for section title matching (default: False)
            - False: "introduction" matches "Introduction"
            - True: Exact case matching required
            - Affects section_title parameter matching
        
        output_format (str): Return data format:
            - "text": Human-readable formatted text (default)
            - "json": Structured JSON object for programmatic processing
            - JSON includes hierarchy, indices, and metadata
        
        include_formatting (bool): Whether to include formatting information (default: False)
            - False: Plain text content only
            - True: Include font, style, and formatting details
            - Adds significant detail to output
        
        formatting_detail (str): Level of formatting information when include_formatting=True:
            - "basic": Font name, size, bold/italic status
            - "detailed": Basic + color, alignment, spacing
            - "comprehensive": Detailed + advanced formatting properties
    
    Returns:
        str: Document structure or content in requested format:
            - mode="overview" + output_format="text": Hierarchical text outline
            - mode="overview" + output_format="json": Structured hierarchy object
            - mode="content" + output_format="text": Section content as text
            - mode="content" + output_format="json": Content with metadata
            - Error message string if operation fails
    
    Use Cases:
        📄 Document Analysis: Understand document structure and organization
        🔍 Content Navigation: Find and extract specific sections quickly
        📚 Academic Review: Analyze paper structure and section content
        📝 Editorial Review: Review document organization and flow
        📊 Content Audit: Assess completeness of structured documents
        🗂️ Section Management: Extract sections for reorganization or reuse
    
    Examples:
        # Basic document structure overview (session-based)
        structure = await get_sections(document_id="main")
        # Returns: Text outline showing all sections up to level 3
        
        # Detailed structure with full content as JSON (legacy filename)
        detailed = await get_sections(filename="thesis.docx", mode="overview", 
                                     full_content=True, output_format="json",
                                     include_formatting=True)
        # Returns: JSON with complete structure and formatted content
        
        # Extract specific methodology section (session-based)
        methods = await get_sections(document_id="draft", mode="content", 
                                   section_title="Methodology",
                                   include_subsections=True)
        # Returns: Full text content of methodology section and subsections
        
        # Get introduction without subsections, formatted (legacy filename)
        intro = await get_sections(filename="manuscript.docx", mode="content",
                                  section_title="Introduction", 
                                  include_subsections=False,
                                  include_formatting=True, formatting_detail="detailed")
        # Returns: Introduction content with detailed formatting
        
        # Overview of main sections only (session-based)
        overview = await get_sections(document_id="report", mode="overview",
                                     max_level=1, output_format="json")
        # Returns: JSON with main chapter/section structure
        
        # Case-sensitive search for specific subsection (session-based)
        subsection = await get_sections(document_id="analysis", mode="content",
                                       section_title="3.2.1 Data Collection",
                                       case_sensitive=True, include_formatting=True)
        # Returns: Exact subsection content with formatting
        
        # Complete document analysis with comprehensive formatting (legacy filename)
        complete = await get_sections(filename="dissertation.docx", mode="overview",
                                     max_level=9, full_content=True,
                                     output_format="json", include_formatting=True,
                                     formatting_detail="comprehensive")
        # Returns: Complete document structure with all formatting details
        
        # Extract conclusion section for review (session-based)
        conclusion = await get_sections(document_id="paper", mode="content",
                                       section_title="Conclusion",
                                       output_format="json")
        # Returns: JSON with conclusion content and metadata
    
    Error Handling:
        - Session errors: "Unable to resolve document from session: {details}"
        - File not found: "Document {filename} does not exist"
        - Invalid mode: "Invalid mode: {mode}. Must be one of: overview, content"
        - Invalid max_level: "Invalid max_level: {level}. Must be between 1-9"
        - Invalid output_format: "Invalid output_format: {format}. Must be one of: text, json"
        - Invalid formatting_detail: "Invalid formatting_detail: {detail}. Must be one of: basic, detailed, comprehensive"
        - Section not found: "Section '{section_title}' not found in document"
        - No headings found: "No heading structure found in document"
        - Document corruption: "Error analyzing document structure: {error_details}"
        - Permission issues: "Cannot read document: {permission_error}"
    
    Academic Writing Workflow:
        1. Structure Analysis: Use overview mode to understand document organization
        2. Section Review: Use content mode to review specific sections
        3. Content Extraction: Extract sections for revision or citation
        4. Quality Check: Verify all required sections are present
        5. Format Review: Use formatting options to check style consistency
        6. Navigation: Use section extraction for quick content location
    
    Professional Document Management:
        - Report Analysis: Review report structure and completeness
        - Content Audit: Ensure all required sections are included
        - Template Compliance: Verify document follows required structure
        - Section Extraction: Extract content for reuse in other documents
        - Quality Assurance: Review section organization and flow
    
    Performance Notes:
        - Large documents with many sections may take longer to process
        - Full content mode is slower than overview mode
        - Formatting extraction adds processing time proportional to detail level
        - JSON output requires additional processing time
        - Consider using max_level to limit scope for large documents
    
    Integration Tips:
        - Use with get_text for detailed paragraph-level analysis
        - Combine with add_text_content for section-based editing
        - Use output data for navigation in other document operations
        - JSON output ideal for programmatic document processing
        - Text output better for human review and analysis
    """
    # Resolve document path from session or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    filename = ensure_docx_extension(filename)
    
    # Validate mode parameter
    valid_modes = ["overview", "content"]
    if mode not in valid_modes:
        return f"Invalid mode: {mode}. Must be one of: {', '.join(valid_modes)}"
    
    # Validate output_format parameter
    valid_formats = ["text", "json"]
    if output_format not in valid_formats:
        return f"Invalid output_format: {output_format}. Must be one of: {', '.join(valid_formats)}"
    
    # Validate formatting_detail parameter
    valid_details = ["basic", "detailed", "comprehensive"]
    if formatting_detail not in valid_details:
        return f"Invalid formatting_detail: {formatting_detail}. Must be one of: {', '.join(valid_details)}"
    
    # Validate max_level parameter
    try:
        max_level = int(max_level)
        if max_level < 1 or max_level > 9:
            return f"Invalid max_level: {max_level}. Must be between 1 and 9."
    except (ValueError, TypeError):
        return "Invalid parameter: max_level must be an integer between 1 and 9"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    def extract_run_formatting(run, detail_level="basic"):
        """Extract formatting information from a run."""
        formatting = {
            "text": run.text,
        }
        
        if detail_level in ["basic", "detailed", "comprehensive"]:
            # Basic formatting
            formatting.update({
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
            })
        
        if detail_level in ["detailed", "comprehensive"]:
            # Detailed formatting
            formatting.update({
                "font_name": run.font.name,
                "font_size": str(run.font.size) if run.font.size else None,
                "font_color": str(run.font.color.rgb) if run.font.color.rgb else None,
                "highlight_color": str(run.font.highlight_color) if run.font.highlight_color else None,
                "strike": run.font.strike,
                "double_strike": run.font.double_strike,
                "superscript": run.font.superscript,
                "subscript": run.font.subscript,
                "small_caps": run.font.small_caps,
                "all_caps": run.font.all_caps,
            })
        
        if detail_level == "comprehensive":
            # Comprehensive formatting
            formatting.update({
                "font_color_theme": str(run.font.color.theme_color) if run.font.color.theme_color else None,
                "emboss": run.font.emboss,
                "imprint": run.font.imprint,
                "outline": run.font.outline,
                "shadow": run.font.shadow,
                "snap_to_grid": run.font.snap_to_grid,
            })
        
        # Clean up None values for cleaner output
        return {k: v for k, v in formatting.items() if v is not None}
    
    def extract_paragraph_formatting(paragraph, detail_level="basic"):
        """Extract formatting information from a paragraph."""
        formatting = {}
        
        if detail_level in ["basic", "detailed", "comprehensive"]:
            # Basic paragraph formatting
            formatting.update({
                "style": paragraph.style.name if paragraph.style else None,
                "alignment": str(paragraph.alignment) if paragraph.alignment else None,
            })
        
        if detail_level in ["detailed", "comprehensive"]:
            # Detailed paragraph formatting
            paragraph_format = paragraph.paragraph_format
            formatting.update({
                "left_indent": str(paragraph_format.left_indent) if paragraph_format.left_indent else None,
                "right_indent": str(paragraph_format.right_indent) if paragraph_format.right_indent else None,
                "first_line_indent": str(paragraph_format.first_line_indent) if paragraph_format.first_line_indent else None,
                "space_before": str(paragraph_format.space_before) if paragraph_format.space_before else None,
                "space_after": str(paragraph_format.space_after) if paragraph_format.space_after else None,
                "line_spacing": str(paragraph_format.line_spacing) if paragraph_format.line_spacing else None,
            })
        
        # Clean up None values for cleaner output
        return {k: v for k, v in formatting.items() if v is not None}
    
    try:
        doc = Document(filename)
        paragraphs = doc.paragraphs
        
        if not paragraphs:
            return "Document contains no paragraphs"
        
        # Extract section information
        sections = []
        current_section = None
        
        for i, paragraph in enumerate(paragraphs):
            # Check if paragraph is a heading
            heading_level = None
            if paragraph.style and paragraph.style.name.startswith('Heading'):
                try:
                    heading_level = int(paragraph.style.name.replace('Heading ', ''))
                except ValueError:
                    continue
            
            if heading_level and heading_level <= max_level:
                # This is a heading - start new section or subsection
                section_info = {
                    "title": paragraph.text.strip(),
                    "level": heading_level,
                    "paragraph_index": i,
                    "content": [],
                    "subsections": []
                }
                
                # Add formatting information if requested
                if include_formatting:
                    section_info["heading_formatting"] = {
                        "paragraph_formatting": extract_paragraph_formatting(paragraph, formatting_detail),
                        "runs": []
                    }
                    
                    for run in paragraph.runs:
                        if run.text.strip():
                            section_info["heading_formatting"]["runs"].append(
                                extract_run_formatting(run, formatting_detail)
                            )
                
                # Check if this matches target section (if specified)
                title_match = False
                if section_title:
                    if case_sensitive:
                        title_match = section_title in paragraph.text
                    else:
                        title_match = section_title.lower() in paragraph.text.lower()
                
                # Add to appropriate location
                if heading_level == 1 or not sections:
                    sections.append(section_info)
                    current_section = section_info
                else:
                    # Find appropriate parent section
                    parent_section = sections[-1]
                    for j in range(len(sections) - 1, -1, -1):
                        if sections[j]["level"] < heading_level:
                            parent_section = sections[j]
                            break
                    
                    if include_subsections:
                        parent_section["subsections"].append(section_info)
                    else:
                        sections.append(section_info)
                    current_section = section_info
            
            elif current_section:
                # This is content - add to current section
                content_text = paragraph.text.strip()
                if content_text:
                    content_item = {
                        "paragraph_index": i,
                        "text": content_text
                    }
                    
                    # Add formatting information if requested
                    if include_formatting:
                        content_item["formatting"] = {
                            "paragraph_formatting": extract_paragraph_formatting(paragraph, formatting_detail),
                            "runs": []
                        }
                        
                        for run in paragraph.runs:
                            if run.text.strip():
                                content_item["formatting"]["runs"].append(
                                    extract_run_formatting(run, formatting_detail)
                                )
                    
                    current_section["content"].append(content_item)
        
        if not sections:
            return "No heading sections found in document. Document may not use heading styles."
        
        # Filter by section_title if specified
        if section_title:
            filtered_sections = []
            for section in sections:
                title_match = False
                if case_sensitive:
                    title_match = section_title in section["title"]
                else:
                    title_match = section_title.lower() in section["title"].lower()
                
                if title_match:
                    filtered_sections.append(section)
            
            if not filtered_sections:
                return f"Section '{section_title}' not found in document"
            
            sections = filtered_sections
        
        # Format output based on mode and output_format
        if output_format == "json":
            result = {
                "mode": mode,
                "include_formatting": include_formatting,
                "formatting_detail": formatting_detail if include_formatting else None,
                "sections": sections
            }
            return json.dumps(result, indent=2)
        
        # Text output formatting
        result_lines = []
        
        if mode == "overview":
            # Structure overview (replaces extract_sections_by_heading)
            result_lines.append("=== DOCUMENT STRUCTURE ===\\n")
            
            def format_section(section, indent_level=0):
                indent = "  " * indent_level
                level_marker = "#" * section["level"]
                
                content_count = len(section["content"])
                content_preview = ""
                
                if section["content"]:
                    if full_content:
                        content_preview = "\\n".join([c["text"] for c in section["content"]])
                    else:
                        first_content = section["content"][0]["text"]
                        content_preview = first_content[:100] + "..." if len(first_content) > 100 else first_content
                
                result_lines.append(f"{indent}{level_marker} {section['title']} [Para {section['paragraph_index']}]")
                
                # Add formatting information if requested
                if include_formatting and "heading_formatting" in section:
                    heading_fmt = section["heading_formatting"]
                    if heading_fmt["paragraph_formatting"]:
                        fmt_info = ", ".join([f"{k}: {v}" for k, v in heading_fmt["paragraph_formatting"].items()])
                        result_lines.append(f"{indent}   Heading Format: {fmt_info}")
                    
                    if heading_fmt["runs"]:
                        for run in heading_fmt["runs"]:
                            run_info = ", ".join([f"{k}: {v}" for k, v in run.items() if k != "text"])
                            if run_info:
                                result_lines.append(f"{indent}   Run Format: {run_info}")
                
                if content_preview:
                    result_lines.append(f"{indent}   Content ({content_count} paragraphs): {content_preview}")
                
                # Process subsections
                for subsection in section["subsections"]:
                    format_section(subsection, indent_level + 1)
            
            for section in sections:
                format_section(section)
        
        else:  # mode == "content"
            # Content extraction (replaces extract_section_content)
            for section in sections:
                def extract_section_content(section):
                    result_lines.append(f"=== {section['title']} ===\\n")
                    
                    # Add heading formatting if requested
                    if include_formatting and "heading_formatting" in section:
                        heading_fmt = section["heading_formatting"]
                        result_lines.append("HEADING FORMATTING:")
                        if heading_fmt["paragraph_formatting"]:
                            for k, v in heading_fmt["paragraph_formatting"].items():
                                result_lines.append(f"  {k}: {v}")
                        
                        for i, run in enumerate(heading_fmt["runs"]):
                            result_lines.append(f"  Run {i+1}:")
                            for k, v in run.items():
                                result_lines.append(f"    {k}: {v}")
                        result_lines.append("")
                    
                    # Add content
                    for content_item in section["content"]:
                        result_lines.append(content_item["text"])
                        
                        # Add content formatting if requested
                        if include_formatting and "formatting" in content_item:
                            content_fmt = content_item["formatting"]
                            result_lines.append("  FORMATTING:")
                            if content_fmt["paragraph_formatting"]:
                                for k, v in content_fmt["paragraph_formatting"].items():
                                    result_lines.append(f"    {k}: {v}")
                            
                            for i, run in enumerate(content_fmt["runs"]):
                                result_lines.append(f"    Run {i+1}:")
                                for k, v in run.items():
                                    result_lines.append(f"      {k}: {v}")
                            result_lines.append("")
                    
                    # Add subsections if enabled
                    if include_subsections:
                        for subsection in section["subsections"]:
                            extract_section_content(subsection)
                
                extract_section_content(section)
        
        return "\\n".join(result_lines)
    
    except Exception as e:
        return f"Failed to extract sections: {str(e)}"


async def generate_table_of_contents(document_id: str = None, filename: str = None, max_level: int = 3, update_existing: bool = True) -> str:
    """Generate a table of contents based on document headings.
    
    Args:
        document_id (str, optional): Session document identifier (preferred)
        filename (str, optional): Path to the Word document
        max_level: Maximum heading level to include (1-9, default 3)
        update_existing: Whether to update existing ToC or create new one (default True)
    
    Returns:
        Success message with ToC information
    """
    # Resolve document path from session or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Collect all headings
        headings = []
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.style and paragraph.style.name.startswith('Heading '):
                try:
                    level = int(paragraph.style.name.split(' ')[1])
                    if level <= max_level:
                        headings.append({
                            'level': level,
                            'text': paragraph.text.strip(),
                            'index': i
                        })
                except (ValueError, IndexError):
                    pass
        
        if not headings:
            return f"No headings found in {filename}. Cannot generate table of contents."
        
        # Find existing ToC or create new one
        toc_inserted = False
        
        # Look for existing "Table of Contents" or "Contents" heading
        for i, paragraph in enumerate(doc.paragraphs):
            if (paragraph.text.lower().strip() in ['table of contents', 'contents', 'toc'] or
                'table of contents' in paragraph.text.lower()):
                
                if update_existing:
                    # Remove existing ToC content (next few paragraphs that aren't headings)
                    j = i + 1
                    to_remove = []
                    while j < len(doc.paragraphs):
                        next_para = doc.paragraphs[j]
                        if (next_para.style and 
                            (next_para.style.name.startswith('Heading ') or
                             next_para.style.name == 'Normal')):
                            # Check if it looks like ToC content
                            if any(h['text'] in next_para.text for h in headings):
                                to_remove.append(j)
                                j += 1
                            else:
                                break
                        else:
                            break
                    
                    # Remove old ToC entries
                    for idx in reversed(to_remove):
                        p = doc.paragraphs[idx]._p
                        p.getparent().remove(p)
                
                # Insert new ToC after the ToC heading
                toc_para = doc.paragraphs[i]
                for heading in headings:
                    # Create ToC entry
                    indent = "    " * (heading['level'] - 1)
                    toc_text = f"{indent}{heading['text']}"
                    
                    # Insert paragraph after ToC heading
                    new_para = doc.add_paragraph(toc_text)
                    # Move the paragraph to correct position
                    new_para._p.getparent().remove(new_para._p)
                    toc_para._p.getparent().insert(
                        list(toc_para._p.getparent()).index(toc_para._p) + 1,
                        new_para._p
                    )
                
                toc_inserted = True
                break
        
        # If no existing ToC found, create one at the beginning
        if not toc_inserted:
            # Insert ToC at the beginning (after any title)
            toc_heading = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph()
            toc_heading.text = "Table of Contents"
            toc_heading.style = doc.styles['Heading 1']
            
            # Add ToC entries
            for heading in headings:
                indent = "    " * (heading['level'] - 1)
                toc_text = f"{indent}{heading['text']}"
                toc_para = doc.add_paragraph(toc_text)
                
                # Move to correct position (after ToC heading)
                toc_para._p.getparent().remove(toc_para._p)
                toc_heading._p.getparent().insert(
                    list(toc_heading._p.getparent()).index(toc_heading._p) + 1,
                    toc_para._p
                )
            
            # Add page break after ToC
            doc.add_page_break()
        
        doc.save(filename)
        
        return f"Table of contents {'updated' if update_existing else 'created'} with {len(headings)} entries (max level {max_level})."
    
    except Exception as e:
        return f"Failed to generate table of contents: {str(e)}"
