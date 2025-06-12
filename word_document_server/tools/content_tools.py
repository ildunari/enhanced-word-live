"""
Content tools for Word Document Server.

These tools add various types of content to Word documents,
including headings, paragraphs, tables, images, and page breaks.
"""
import os
from typing import List, Optional, Dict, Any
from docx import Document
from docx.text.run import Run
from docx.shared import Inches, Pt

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension, validate_docx_path
from word_document_server.utils.document_utils import find_and_replace_text
from word_document_server.utils.session_utils import resolve_document_path
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def add_text_content(
    document_id: str = None,
    filename: str = None,
    text: str = None,
    content_type: str = "paragraph",
    level: Optional[int] = None,
    style: Optional[str] = None,
    position: str = "end",
    insert_before_paragraph: Optional[int] = None,
    insert_after_paragraph: Optional[int] = None
) -> str:
    """Unified text content addition function for comprehensive document content management.
    
    This function consolidates paragraph and heading creation into a single comprehensive tool.
    It replaces add_paragraph and add_heading with enhanced positioning, styling, and
    formatting options for professional document creation and editing workflows.
    
    Args:
        document_id (str): Session document ID (preferred)
        filename (str): Path to the Word document (legacy, for backward compatibility)
        
        text (str): Content text to add to the document
            - Supports plain text with basic formatting preservation
            - Can include special characters and unicode
            - Line breaks will be preserved in paragraph content
            - Example: "This is the introduction paragraph with important concepts."
        
        content_type (str): Type of content element to create:
            - "paragraph": Regular text paragraph (default)
            - "heading": Document heading/title with hierarchical level
            - Paragraphs are for body content, headings for structure
        
        level (int, optional): Heading hierarchy level (required for content_type="heading")
            - 1: Main title/chapter heading (largest)
            - 2: Section heading
            - 3: Subsection heading
            - 4-9: Sub-subsection headings (progressively smaller)
            - Only used when content_type="heading"
        
        style (str, optional): Document style name to apply to content
            - None: Use default paragraph/heading style (default)
            - "Normal": Standard paragraph style
            - "Quote": Indented quotation style
            - "Emphasis": Emphasized text style
            - "Caption": Figure/table caption style
            - Custom: Any style defined in document template
        
        position (str): Placement position within document:
            - "end": Append to end of document (default)
            - "beginning": Insert at document start
            - "before": Insert before specific paragraph (requires insert_before_paragraph)
            - "after": Insert after specific paragraph (requires insert_after_paragraph)
        
        insert_before_paragraph (int, optional): Zero-based paragraph index for position="before"
            - Content will be inserted before this paragraph
            - Existing paragraph indices will shift down
            - Example: 5 inserts before the 6th paragraph
        
        insert_after_paragraph (int, optional): Zero-based paragraph index for position="after"
            - Content will be inserted after this paragraph
            - Subsequent paragraph indices will shift down
            - Example: 3 inserts after the 4th paragraph
    
    Returns:
        str: Status message indicating success or failure:
            - Success: "Successfully added {content_type} at {position}"
            - Error: Specific error message with troubleshooting guidance
    
    Use Cases:
        📝 Document Creation: Build structured documents with headings and content
        📚 Academic Writing: Add chapters, sections, and content paragraphs
        📄 Report Writing: Insert executive summaries, findings, conclusions
        📋 Technical Documentation: Add procedure steps and explanations
        ✏️ Manuscript Editing: Insert new content at specific locations
        📊 Business Documents: Add formatted content with professional styling
    
    Examples:
        # Add simple paragraph at document end
        result = await add_text_content(document_id="report", 
                                       text="This paragraph summarizes the key findings of our research.")
        # Returns: "Successfully added paragraph at end"
        
        # Create main chapter heading
        result = await add_text_content(document_id="thesis", text="Chapter 3: Methodology", 
                                       content_type="heading", level=1)
        # Returns: "Successfully added heading at end"
        
        # Insert section heading with custom style
        result = await add_text_content(document_id="manuscript", text="3.1 Data Collection",
                                       content_type="heading", level=2, style="Heading 2")
        # Returns: "Successfully added heading at end"
        
        # Insert paragraph before specific location
        result = await add_text_content(document_id="document", 
                                       text="This new paragraph provides important context.",
                                       position="before", insert_before_paragraph=5)
        # Returns: "Successfully added paragraph at before"
        
        # Add quotation with quote style
        result = await add_text_content(document_id="essay", 
                                       text="To be or not to be, that is the question.",
                                       style="Quote", position="after", insert_after_paragraph=2)
        # Returns: "Successfully added paragraph at after"
        
        # Insert introduction paragraph at document beginning
        result = await add_text_content(document_id="paper", 
                                       text="This document presents comprehensive analysis of market trends.",
                                       position="beginning")
        # Returns: "Successfully added paragraph at beginning"
        
        # Add subsection heading in middle of document
        result = await add_text_content(document_id="manual", text="2.3.1 Installation Steps",
                                       content_type="heading", level=3,
                                       position="before", insert_before_paragraph=15)
        # Returns: "Successfully added heading at before"
        
        # Add emphasized conclusion paragraph
        result = await add_text_content(document_id="summary",
                                       text="In conclusion, the results demonstrate significant improvements.",
                                       style="Emphasis", position="end")
        # Returns: "Successfully added paragraph at end"
    
    Error Handling:
        - Document not found: "Document '{document_id}' not found in sessions"
        - File not writable: "Cannot modify document: {reason}. Consider creating a copy first."
        - Invalid content_type: "Invalid content_type: {type}. Must be one of: paragraph, heading"
        - Invalid position: "Invalid position: {pos}. Must be one of: end, beginning, before, after"
        - Missing level for heading: "Heading level (1-9) is required for content_type='heading'"
        - Invalid level: "Invalid heading level: {level}. Must be between 1-9"
        - Missing position parameter: "insert_before_paragraph required for position='before'"
        - Invalid paragraph index: "Paragraph index {index} is out of range"
        - Empty text: "Text content cannot be empty"
        - Style not found: "Style '{style}' not found in document template"
        - Document corruption: "Error adding content: {error_details}"
    
    Document Structure Workflow:
        1. Planning: Design document hierarchy with heading levels
        2. Structure Creation: Add main headings (level 1) first
        3. Section Development: Add subsection headings (levels 2-3)
        4. Content Addition: Insert paragraphs under appropriate headings
        5. Refinement: Use positioning options to reorganize content
        6. Styling: Apply consistent styles throughout document
    
    Academic Writing Best Practices:
        - Use level 1 for chapter headings
        - Use level 2 for major section headings
        - Use level 3 for subsection headings
        - Maintain consistent heading hierarchy
        - Insert content paragraphs after relevant headings
        - Use appropriate styles for different content types
    
    Performance Notes:
        - Large documents may take longer for content insertion
        - Position-specific insertions require document traversal
        - Style application adds minimal processing time
        - Batch multiple content additions for efficiency
        - Complex positioning operations may be slower than simple appends
    
    Integration with Other Tools:
        - Use get_sections to understand document structure before insertion
        - Use get_text to verify content placement after insertion
        - Combine with add_note for comprehensive content creation
        - Use manage_protection to control editing permissions
    """
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Resolve document path from document_id or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    # Validate required parameters
    if not text:
        return "Error: text parameter is required"
    
    # Validate content_type parameter
    valid_types = ["paragraph", "heading"]
    if content_type not in valid_types:
        return f"Invalid content_type: {content_type}. Must be one of: {', '.join(valid_types)}"
    
    # Validate position parameter
    valid_positions = ["end", "beginning", "before", "after"]
    if position not in valid_positions:
        return f"Invalid position: {position}. Must be one of: {', '.join(valid_positions)}"
    
    # Validate heading-specific parameters
    if content_type == "heading":
        if level is None:
            return "Invalid parameter: level is required when content_type='heading'"
        
        try:
            level = int(level)
            if level < 1 or level > 9:
                return f"Invalid heading level: {level}. Level must be between 1 and 9."
        except (ValueError, TypeError):
            return "Invalid parameter: level must be an integer between 1 and 9"
    
    # Validate position-specific parameters
    if position == "before" and insert_before_paragraph is None:
        return "Invalid parameter: insert_before_paragraph is required when position='before'"
    
    if position == "after" and insert_after_paragraph is None:
        return "Invalid parameter: insert_after_paragraph is required when position='after'"
    
    # Validate paragraph index parameters
    for param_name, param_value in [("insert_before_paragraph", insert_before_paragraph), 
                                   ("insert_after_paragraph", insert_after_paragraph)]:
        if param_value is not None:
            try:
                param_value = int(param_value)
                if param_value < 0:
                    return f"Invalid parameter: {param_name} must be a non-negative integer"
            except (ValueError, TypeError):
                return f"Invalid parameter: {param_name} must be an integer"
    
    if not text.strip():
        return "Invalid parameter: text cannot be empty"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph indices against document
        if insert_before_paragraph is not None:
            insert_before_paragraph = int(insert_before_paragraph)
            if insert_before_paragraph >= len(doc.paragraphs):
                return f"Invalid insert_before_paragraph: {insert_before_paragraph}. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        if insert_after_paragraph is not None:
            insert_after_paragraph = int(insert_after_paragraph)
            if insert_after_paragraph >= len(doc.paragraphs):
                return f"Invalid insert_after_paragraph: {insert_after_paragraph}. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        # Create the content element
        if content_type == "heading":
            # Ensure heading styles exist
            ensure_heading_style(doc)
            
            try:
                if position in ["end", "beginning"]:
                    # Add at document level
                    if position == "beginning":
                        # Insert at beginning - need to use paragraph insertion
                        new_paragraph = doc.paragraphs[0]._element.getparent().insert(0, doc.add_heading(text, level=level)._element)
                        heading = doc.add_heading(text, level=level)
                    else:
                        heading = doc.add_heading(text, level=level)
                    
                    created_element = heading
                    success_message = f"Heading '{text}' (level {level}) added"
                else:
                    # Insert at specific position
                    heading = doc.add_heading(text, level=level)
                    created_element = heading
                    success_message = f"Heading '{text}' (level {level}) inserted"
                    
            except Exception:
                # Fallback to direct formatting if heading styles fail
                paragraph = doc.add_paragraph(text)
                paragraph.style = doc.styles['Normal']
                run = paragraph.runs[0]
                run.bold = True
                
                # Adjust size based on heading level
                if level == 1:
                    run.font.size = Pt(16)
                elif level == 2:
                    run.font.size = Pt(14)
                else:
                    run.font.size = Pt(12)
                    
                created_element = paragraph
                success_message = f"Heading '{text}' added with direct formatting"
        
        else:  # paragraph
            paragraph = doc.add_paragraph(text)
            created_element = paragraph
            success_message = f"Paragraph added"
            
            # Apply style if specified
            if style:
                try:
                    paragraph.style = style
                except KeyError:
                    paragraph.style = doc.styles['Normal']
                    success_message += f" (style '{style}' not found, used default)"
        
        # Handle positioning for before/after insertions
        if position == "before":
            # Move element to before specified paragraph
            target_paragraph = doc.paragraphs[insert_before_paragraph]
            target_paragraph._element.getparent().insert(
                list(target_paragraph._element.getparent()).index(target_paragraph._element),
                created_element._element
            )
            success_message += f" before paragraph {insert_before_paragraph}"
        
        elif position == "after":
            # Move element to after specified paragraph
            target_paragraph = doc.paragraphs[insert_after_paragraph]
            target_paragraph._element.getparent().insert(
                list(target_paragraph._element.getparent()).index(target_paragraph._element) + 1,
                created_element._element
            )
            success_message += f" after paragraph {insert_after_paragraph}"
        
        elif position == "beginning" and content_type == "paragraph":
            # Move paragraph to beginning
            doc._body._element.insert(0, created_element._element)
            success_message += " at document beginning"
        
        doc.save(filename)
        return f"{success_message} to {filename}"
    
    except Exception as e:
        return f"Failed to add {content_type}: {str(e)}"


async def add_table(document_id: str = None, filename: str = None, rows: int = None, cols: int = None, data: Optional[List[List[str]]] = None) -> str:
    """Add a table to a Word document.
    
    Args:
        document_id (str, optional): Session document identifier (preferred)
        filename (str, optional): Path to the Word document
        rows: Number of rows in the table
        cols: Number of columns in the table
        data: Optional 2D array of data to fill the table
    """
    # Resolve document path from session or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        # Suggest creating a copy
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(filename)
        table = doc.add_table(rows=rows, cols=cols)
        
        # Try to set the table style
        try:
            table.style = 'Table Grid'
        except KeyError:
            # If style doesn't exist, add basic borders
            pass
        
        # Fill table with data if provided
        if data:
            for i, row_data in enumerate(data):
                if i >= rows:
                    break
                for j, cell_text in enumerate(row_data):
                    if j >= cols:
                        break
                    table.cell(i, j).text = str(cell_text)
        
        doc.save(filename)
        return f"Table ({rows}x{cols}) added to {filename}"
    except Exception as e:
        return f"Failed to add table: {str(e)}"


async def add_picture(document_id: str = None, filename: str = None, image_path: str = None, width: Optional[float] = None) -> str:
    """Add an image to a Word document.
    
    Args:
        document_id (str, optional): Session document identifier (preferred)
        filename (str, optional): Path to the Word document
        image_path: Path to the image file
        width: Optional width in inches (proportional scaling)
    """
    # Resolve document path from session or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    filename = ensure_docx_extension(filename)
    
    # Validate document existence
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Get absolute paths for better diagnostics
    abs_filename = os.path.abspath(filename)
    abs_image_path = os.path.abspath(image_path)
    
    # Validate image existence with improved error message
    if not os.path.exists(abs_image_path):
        return f"Image file not found: {abs_image_path}"
    
    # Check image file size
    try:
        image_size = os.path.getsize(abs_image_path) / 1024  # Size in KB
        if image_size <= 0:
            return f"Image file appears to be empty: {abs_image_path} (0 KB)"
    except Exception as size_error:
        return f"Error checking image file: {str(size_error)}"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(abs_filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first or creating a new document."
    
    try:
        doc = Document(abs_filename)
        # Additional diagnostic info
        diagnostic = f"Attempting to add image ({abs_image_path}, {image_size:.2f} KB) to document ({abs_filename})"
        
        try:
            if width:
                doc.add_picture(abs_image_path, width=Inches(width))
            else:
                doc.add_picture(abs_image_path)
            doc.save(abs_filename)
            return f"Picture {image_path} added to {filename}"
        except Exception as inner_error:
            # More detailed error for the specific operation
            error_type = type(inner_error).__name__
            error_msg = str(inner_error)
            return f"Failed to add picture: {error_type} - {error_msg or 'No error details available'}\nDiagnostic info: {diagnostic}"
    except Exception as outer_error:
        # Fallback error handling
        error_type = type(outer_error).__name__
        error_msg = str(outer_error)
        return f"Document processing error: {error_type} - {error_msg or 'No error details available'}"





def enhanced_search_and_replace(document_id: str = None, filename: str = None, 
                                    find_text: str = None, replace_text: str = None,
                                    apply_formatting: bool = False,
                                    bold: Optional[bool] = None, 
                                    italic: Optional[bool] = None,
                                    underline: Optional[bool] = None, 
                                    color: Optional[str] = None,
                                    font_size: Optional[int] = None, 
                                    font_name: Optional[str] = None,
                                    match_case: bool = True,
                                    whole_words_only: bool = False,
                                    use_regex: bool = False) -> str:
    """Enhanced search and replace with formatting options, regex support, and case-insensitive matching.
    
    Provides powerful text replacement capabilities with:
    - Regex pattern matching
    - Case-sensitive/insensitive search
    - Whole word matching
    - Advanced formatting application to replaced text
    - Table content support
    
    Args:
        document_id: Session document ID (preferred)
        filename: Path to the Word document (legacy, for backward compatibility)
        find_text: Text or regex pattern to search for
        replace_text: Text to replace with (supports regex groups like $1, $2 if use_regex=True)
        apply_formatting: Whether to apply formatting to the replaced text
        bold: Set replaced text bold (True/False)
        italic: Set replaced text italic (True/False)
        underline: Set replaced text underlined (True/False)
        color: Text color for replaced text (e.g., 'red', 'blue', '#FF0000')
        font_size: Font size in points for replaced text
        font_name: Font name/family for replaced text
        match_case: Whether to match case exactly (default True, ignored if use_regex=True)
        whole_words_only: Whether to match whole words only (default False)
        use_regex: Enable regex pattern matching (default False)
    
    Returns:
        Status message with replacement count and details
        
    Examples:
        # Simple replacement
        enhanced_search_and_replace(document_id="main", find_text="old text", replace_text="new text")
        
        # Case-insensitive replacement with formatting
        enhanced_search_and_replace(document_id="main", find_text="Important", replace_text="CRITICAL", 
                                   match_case=False, apply_formatting=True, 
                                   bold=True, color="red")
        
        # Regex pattern replacement
        enhanced_search_and_replace(document_id="main", find_text=r"(\\\\d{4})-(\\\\d{2})-(\\\\d{2})", 
                                   replace_text=r"$2/$3/$1", use_regex=True)
                                   
        # Whole word replacement with font styling
        enhanced_search_and_replace(document_id="main", find_text="AI", replace_text="Artificial Intelligence", 
                                   whole_words_only=True, apply_formatting=True,
                                   font_name="Arial", font_size=12)
    """
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Resolve document path from document_id or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    # Validate required parameters
    if not find_text:
        return "Error: find_text parameter is required"
    
    if not replace_text:
        return "Error: replace_text parameter is required"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    # Validate regex pattern if using regex
    if use_regex:
        try:
            import re
            re.compile(find_text)
        except re.error as e:
            return f"Invalid regex pattern '{find_text}': {str(e)}"
    
    try:
        doc = Document(filename)
        
        count = _enhanced_replace_in_paragraphs(doc.paragraphs, find_text, replace_text, 
                                              apply_formatting, bold, italic, underline, 
                                              color, font_size, font_name, match_case, 
                                              whole_words_only, use_regex)
        
        # Search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    count += _enhanced_replace_in_paragraphs(cell.paragraphs, find_text, replace_text,
                                                           apply_formatting, bold, italic, underline,
                                                           color, font_size, font_name, match_case, 
                                                           whole_words_only, use_regex)
        
        if count > 0:
            doc.save(filename)
            search_type = "regex pattern" if use_regex else "text"
            case_info = "" if match_case else " (case-insensitive)"
            word_info = " (whole words only)" if whole_words_only else ""
            formatting_applied = " with formatting" if apply_formatting else ""
            
            return f"Replaced {count} occurrence(s) of {search_type} '{find_text}'{case_info}{word_info} with '{replace_text}'{formatting_applied}."
        else:
            search_type = "regex pattern" if use_regex else "text"
            return f"No occurrences of {search_type} '{find_text}' found."
            
    except FileNotFoundError:
        return f"Document {filename} not found"
    except PermissionError:
        return f"Permission denied accessing {filename}"
    except Exception as e:
        return f"Failed to search and replace: {str(e)}"




def _enhanced_replace_in_paragraphs(paragraphs, find_text, replace_text, apply_formatting,
                                   bold, italic, underline, color, font_size, font_name,
                                   match_case, whole_words_only, use_regex=False):
    """Helper function to replace text in paragraphs with optional formatting and regex support.
    
    This implementation fixes the formatting over-application bug by creating 
    new runs for replaced text instead of applying formatting to entire existing runs.
    Supports both literal text matching and regex pattern matching.
    """
    import re
    from docx.shared import Pt, RGBColor
    
    count = 0
    
    for para in paragraphs:
        para_text = para.text
        
        # Create search pattern based on options
        if use_regex:
            try:
                pattern = find_text
                flags = re.IGNORECASE if not match_case else 0
            except re.error:
                continue  # Skip invalid regex patterns
        else:
            # Escape special regex characters for literal matching
            escaped_text = re.escape(find_text)
            if whole_words_only:
                pattern = r'\b' + escaped_text + r'\b'
            else:
                pattern = escaped_text
            flags = re.IGNORECASE if not match_case else 0
        
        # Find all matches in the paragraph text
        try:
            matches = list(re.finditer(pattern, para_text, flags))
        except re.error:
            continue  # Skip if pattern compilation fails
            
        if not matches:
            continue
            
        count += len(matches)
        
        # Process matches from right to left to avoid position shifting
        for match in reversed(matches):
            start_pos = match.start()
            end_pos = match.end()
            
            # For regex, get the actual replacement text (may include group substitutions)
            if use_regex:
                actual_replace_text = match.expand(replace_text)
            else:
                actual_replace_text = replace_text
            
            # Find which runs contain this text span
            current_pos = 0
            start_run_idx = None
            end_run_idx = None
            start_run_offset = 0
            end_run_offset = 0
            
            for run_idx, run in enumerate(para.runs):
                run_length = len(run.text)
                
                # Check if this run contains the start position
                if start_run_idx is None and current_pos <= start_pos < current_pos + run_length:
                    start_run_idx = run_idx
                    start_run_offset = start_pos - current_pos
                
                # Check if this run contains the end position
                if current_pos < end_pos <= current_pos + run_length:
                    end_run_idx = run_idx
                    end_run_offset = end_pos - current_pos
                    break
                    
                current_pos += run_length
            
            if start_run_idx is None or end_run_idx is None:
                continue  # Skip if we can't locate the text properly
            
            # Replace text across the identified runs
            if start_run_idx == end_run_idx:
                # Text is within a single run - split it into three parts
                run = para.runs[start_run_idx]
                old_text = run.text
                
                # Split: before_text + replaced_text + after_text
                before_text = old_text[:start_run_offset]
                after_text = old_text[end_run_offset:]
                
                # Update original run with before_text
                run.text = before_text
                
                # Create new run for replaced text with optional formatting
                new_run = para.add_run(actual_replace_text)
                if apply_formatting:
                    _copy_run_formatting(run, new_run)  # Copy base formatting first
                    _apply_formatting_to_run(new_run, bold, italic, underline, color, font_size, font_name)
                else:
                    _copy_run_formatting(run, new_run)  # Preserve original formatting
                
                # Create new run for after_text if there is any
                if after_text:
                    after_run = para.add_run(after_text)
                    _copy_run_formatting(run, after_run)
                    
            else:
                # Text spans multiple runs - more complex replacement
                # Remove the matched text from all affected runs
                for run_idx in range(start_run_idx, end_run_idx + 1):
                    run = para.runs[run_idx]
                    if run_idx == start_run_idx:
                        # First run: keep text before match
                        run.text = run.text[:start_run_offset]
                    elif run_idx == end_run_idx:
                        # Last run: keep text after match
                        run.text = run.text[end_run_offset:]
                    else:
                        # Middle runs: clear completely
                        run.text = ""
                
                # Add the replacement text as a new run after the first affected run
                new_run = para.add_run(actual_replace_text)
                if apply_formatting:
                    _copy_run_formatting(para.runs[start_run_idx], new_run)
                    _apply_formatting_to_run(new_run, bold, italic, underline, color, font_size, font_name)
                else:
                    _copy_run_formatting(para.runs[start_run_idx], new_run)
    
    return count




def _apply_formatting_to_run(run, bold, italic, underline, color, font_size, font_name):
    """Apply formatting to a run with error handling."""
    try:
        if bold is not None:
            run.bold = bold
        if italic is not None:
            run.italic = italic
        if underline is not None:
            run.underline = underline
        if color:
            _apply_color_to_run(run, color)
        if font_size:
            from docx.shared import Pt
            run.font.size = Pt(font_size)
        if font_name:
            run.font.name = font_name
    except Exception:
        # Silently continue if formatting fails
        pass

def _copy_run_formatting(source_run, target_run):
    """Copy formatting from source run to target run."""
    try:
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb
    except Exception:
        # Silently continue if copying formatting fails
        pass


def _apply_color_to_run(run, color):
    """Apply color to a run with error handling."""
    from docx.shared import RGBColor
    
    # Define common RGB colors
    color_map = {
        'red': RGBColor(255, 0, 0),
        'blue': RGBColor(0, 0, 255),
        'green': RGBColor(0, 128, 0),
        'yellow': RGBColor(255, 255, 0),
        'black': RGBColor(0, 0, 0),
        'gray': RGBColor(128, 128, 128),
        'white': RGBColor(255, 255, 255),
        'purple': RGBColor(128, 0, 128),
        'orange': RGBColor(255, 165, 0),
        'brown': RGBColor(165, 42, 42),
        'pink': RGBColor(255, 192, 203),
        'cyan': RGBColor(0, 255, 255),
        'magenta': RGBColor(255, 0, 255),
        'lime': RGBColor(0, 255, 0),
        'navy': RGBColor(0, 0, 128),
        'maroon': RGBColor(128, 0, 0),
        'olive': RGBColor(128, 128, 0),
        'teal': RGBColor(0, 128, 128)
    }
    
    try:
        if color.lower() in color_map:
            run.font.color.rgb = color_map[color.lower()]
        else:
            # Try to parse as hex color (e.g., "#FF0000")
            if color.startswith('#') and len(color) == 7:
                hex_color = color[1:]
                r = int(hex_color[0:2], 16)
                g = int(hex_color[2:4], 16)
                b = int(hex_color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
            else:
                # Default to black if color not recognized
                run.font.color.rgb = RGBColor(0, 0, 0)
    except Exception:
        # Fallback to black
        run.font.color.rgb = RGBColor(0, 0, 0)


def format_specific_words(filename: str, word_list: List[str], 
                               bold: Optional[bool] = None,
                               italic: Optional[bool] = None,
                               underline: Optional[bool] = None,
                               color: Optional[str] = None,
                               font_size: Optional[int] = None,
                               font_name: Optional[str] = None,
                               match_case: bool = True,
                               whole_words_only: bool = True) -> str:
    """Format specific words throughout the document using enhanced search and replace.
    
    Args:
        filename: Path to the Word document
        word_list: List of words to format
        bold: Set text bold (True/False)
        italic: Set text italic (True/False)
        underline: Set text underlined (True/False)
        color: Text color (e.g., 'red', 'blue', etc.)
        font_size: Font size in points
        font_name: Font name/family
        match_case: Whether to match case (default True)
        whole_words_only: Whether to match whole words only (default True)
    """
    total_count = 0
    results = []
    
    for word in word_list:
        # Use enhanced search and replace with same text for find and replace
        result = enhanced_search_and_replace(
            filename=filename,
            find_text=word,
            replace_text=word,  # Same text, just apply formatting
            apply_formatting=True,
            bold=bold,
            italic=italic,
            underline=underline,
            color=color,
            font_size=font_size,
            font_name=font_name,
            match_case=match_case,
            whole_words_only=whole_words_only
        )
        results.append(f"'{word}': {result}")
    
    return "\n".join(results)


def format_research_paper_terms(filename: str) -> str:
    """Format common research terms in a PCL paper with appropriate styling - Academic research helper."""
    
    # Format drug names in blue and bold
    drug_names = ["dolutegravir", "meloxicam", "dexamethasone", "DTG", "MLX", "DEX"]
    format_specific_words(filename, drug_names, bold=True, color="blue")
    
    # Format polymer terms in green
    polymer_terms = ["polycaprolactone", "PCL", "mesophase", "crystallinity"]
    format_specific_words(filename, polymer_terms, color="green")
    
    # Format statistical terms in red and italic
    stats_terms = ["p < 0.05", "significant", "correlation", "ANOVA"]
    format_specific_words(filename, stats_terms, italic=True, color="red")
    
    # Format temperature values in orange
    temp_terms = ["25°C", "50°C"]
    format_specific_words(filename, temp_terms, color="orange")
    
    return "Research paper terms formatted successfully!"


def format_document(
    action: str,
    filename: str,
    word_list: Optional[List[str]] = None,
    bold: Optional[bool] = None,
    italic: Optional[bool] = None,
    underline: Optional[bool] = None,
    color: Optional[str] = None,
    font_size: Optional[int] = None,
    font_name: Optional[str] = None,
    match_case: bool = True,
    whole_words_only: bool = True
) -> str:
    """Unified document formatting function for specialized formatting operations.
    
    This consolidated tool replaces 2 individual formatting functions with a single
    action-based interface, reducing tool count while preserving 100% functionality.
    
    Args:
        action (str): Formatting operation to perform:
            - "words": Format specific words in document (requires word_list)
            - "research": Apply research paper formatting (automatic terms)
        filename (str): Path to Word document
        word_list (List[str], optional): List of words to format (required for "words" action)
        bold (bool, optional): Set text bold (True/False)
        italic (bool, optional): Set text italic (True/False)
        underline (bool, optional): Set text underlined (True/False)
        color (str, optional): Text color (e.g., 'red', 'blue', etc.)
        font_size (int, optional): Font size in points
        font_name (str, optional): Font name/family
        match_case (bool): Whether to match case (default True)
        whole_words_only (bool): Whether to match whole words only (default True)
        
    Returns:
        str: Formatting operation result message
        
    Examples:
        # Format specific words with custom styling
        format_document("words", "thesis.docx", 
                       word_list=["important", "critical", "significant"],
                       bold=True, color="red")
        
        # Apply research paper formatting
        format_document("research", "research_paper.docx")
        
        # Format technical terms
        format_document("words", "manual.docx",
                       word_list=["API", "SDK", "REST"],
                       font_name="Courier New", font_size=11)
    """
    # Validate action parameter
    valid_actions = ["words", "research"]
    if action not in valid_actions:
        return f"Invalid action: {action}. Must be one of: {', '.join(valid_actions)}"
    
    # Validate required parameters for each action
    if action == "words" and not word_list:
        return "Error: 'word_list' is required for action 'words'"
    
    # Delegate to appropriate original function based on action
    try:
        if action == "words":
            return format_specific_words(
                filename=filename,
                word_list=word_list,
                bold=bold,
                italic=italic,
                underline=underline,
                color=color,
                font_size=font_size,
                font_name=font_name,
                match_case=match_case,
                whole_words_only=whole_words_only
            )
            
        elif action == "research":
            return format_research_paper_terms(filename)
            
    except Exception as e:
        return f"Error in format_document: {str(e)}"
