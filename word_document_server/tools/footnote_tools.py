"""
Footnote and endnote tools for Word Document Server.

These tools handle footnote and endnote functionality,
including adding, customizing, and converting between them.
"""
import os
from typing import Optional
from docx import Document

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension


async def add_note(
    document_id: str = None,
    filename: str = None,
    paragraph_index: int = None,
    note_text: str = None,
    note_type: str = "footnote",
    position: str = "end",
    symbol: Optional[str] = None
) -> str:
    """Unified note addition function for comprehensive footnote and endnote management.
    
    This function consolidates footnote and endnote creation into a single comprehensive tool.
    It replaces add_footnote_to_document and add_endnote_to_document with enhanced
    positioning and formatting options for academic and professional documentation.
    
    Args:
        document_id (str): Session document ID (preferred)
        filename (str): Path to the Word document (legacy, for backward compatibility)
        
        paragraph_index (int): Zero-based index of target paragraph for note placement
            - 0 = first paragraph, 1 = second paragraph, etc.
            - Must be valid index within document range
            - Note reference will be inserted in this paragraph
        
        note_text (str): Content text for the note
            - Supports plain text and basic formatting
            - Can include citations, references, or explanatory content
            - No length limit but consider readability
            - Example: "See Smith et al. (2023) for detailed methodology"
        
        note_type (str): Type of note to create:
            - "footnote": Note appears at bottom of page (default)
            - "endnote": Note appears at end of document or section
            - Footnotes are better for brief comments or citations
            - Endnotes are better for longer explanations
        
        position (str): Placement position within the target paragraph:
            - "end": Insert reference at end of paragraph (default)
            - "beginning": Insert reference at start of paragraph
            - Affects where the note number/symbol appears in text
        
        symbol (str, optional): Custom symbol for note reference
            - None: Use automatic numbering (1, 2, 3...) (default)
            - "*": Single asterisk
            - "†": Dagger symbol
            - "‡": Double dagger
            - Custom: Any single character or short string
            - Note: Custom symbols may not auto-increment
    
    Returns:
        str: Status message indicating success or failure:
            - Success: "Successfully added {note_type} to paragraph {index}"
            - Error: Specific error message with troubleshooting guidance
    
    Use Cases:
        📚 Academic Writing: Add citations, references, and explanatory notes
        📝 Research Papers: Include methodology notes and data sources
        📄 Legal Documents: Add statutory references and case citations
        📊 Reports: Include data sources and calculation explanations
        ✏️ Manuscripts: Add author notes and editorial comments
        📋 Technical Documentation: Include detailed specifications
    
    Examples:
        # Basic footnote for citation
        result = await add_note(document_id="paper", paragraph_index=5, 
                               note_text="Smith, J. (2023). Advanced Research Methods. Academic Press.",
                               note_type="footnote")
        # Returns: "Successfully added footnote to paragraph 5"
        
        # Endnote for detailed explanation
        result = await add_note(document_id="thesis", paragraph_index=12,
                               note_text="The methodology was adapted from previous studies with modifications for current context.",
                               note_type="endnote", position="end")
        # Returns: "Successfully added endnote to paragraph 12"
        
        # Footnote with custom asterisk symbol
        result = await add_note(document_id="manuscript", paragraph_index=3,
                               note_text="Preliminary results only, full analysis pending.",
                               note_type="footnote", symbol="*")
        # Returns: "Successfully added footnote to paragraph 3"
        
        # Beginning position footnote for emphasis
        result = await add_note(document_id="report", paragraph_index=0,
                               note_text="This section contains confidential information.",
                               note_type="footnote", position="beginning", symbol="†")
        # Returns: "Successfully added footnote to paragraph 0"
        
        # Academic reference endnote
        result = await add_note(document_id="dissertation", paragraph_index=45,
                               note_text="For comprehensive review of this topic, see Johnson et al. (2022), chapters 3-5.",
                               note_type="endnote")
        # Returns: "Successfully added endnote to paragraph 45"
        
        # Legal citation footnote
        result = await add_note(document_id="legal_brief", paragraph_index=8,
                               note_text="See 42 U.S.C. § 1983 (1871) and subsequent amendments.",
                               note_type="footnote", position="end")
        # Returns: "Successfully added footnote to paragraph 8"
        
        # Technical specification note
        result = await add_note(document_id="specification", paragraph_index=15,
                               note_text="Implementation details available in Appendix C, section 2.4.",
                               note_type="endnote", symbol="‡")
        # Returns: "Successfully added endnote to paragraph 15"
    
    Error Handling:
        - Document not found: "Document '{document_id}' not found in sessions"
        - File not writable: "Cannot modify document: {reason}. Consider creating a copy first."
        - Invalid paragraph index: "Paragraph index {index} is out of range (document has {count} paragraphs)"
        - Invalid note_type: "Invalid note_type: {type}. Must be one of: footnote, endnote"
        - Invalid position: "Invalid position: {pos}. Must be one of: end, beginning"
        - Empty note text: "Note text cannot be empty"
        - Document corruption: "Error adding note: {error_details}"
        - Protection conflict: "Document is protected and notes cannot be added"
    
    Academic Workflow Integration:
        1. Content Creation: Write main text content using add_text_content
        2. Reference Addition: Use add_note to insert citations and explanations
        3. Structure Review: Use get_sections to verify note placement
        4. Final Review: Use get_text with search to find and verify all notes
        5. Format Check: Ensure consistent note formatting throughout document
    
    Best Practices:
        - Use footnotes for brief citations and immediate clarifications
        - Use endnotes for longer explanations that might interrupt text flow
        - Maintain consistent note style throughout document
        - Consider using custom symbols sparingly for special emphasis
        - Place notes strategically to support but not overwhelm main text
        - Review note placement in context of paragraph content
    
    Performance Notes:
        - Adding notes to large documents may take longer to process
        - Custom symbols require additional formatting time
        - Documents with many existing notes may slow processing
        - Consider batch operations for multiple notes in same document
    """
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Resolve document path from document_id or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    # Validate required parameters
    if paragraph_index is None:
        return "Error: paragraph_index parameter is required"
    
    if not note_text:
        return "Error: note_text parameter is required"
    
    # Validate note_type parameter
    valid_types = ["footnote", "endnote"]
    if note_type not in valid_types:
        return f"Invalid note_type: {note_type}. Must be one of: {', '.join(valid_types)}"
    
    # Validate position parameter
    valid_positions = ["end", "beginning"]
    if position not in valid_positions:
        return f"Invalid position: {position}. Must be one of: {', '.join(valid_positions)}"
    
    # Ensure paragraph_index is an integer
    try:
        paragraph_index = int(paragraph_index)
        if paragraph_index < 0:
            return "Invalid parameter: paragraph_index must be a non-negative integer"
    except (ValueError, TypeError):
        return "Invalid parameter: paragraph_index must be an integer"
    
    if not note_text.strip():
        return "Invalid parameter: note_text cannot be empty"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot modify document: {error_message}. Consider creating a copy first."
    
    try:
        doc = Document(filename)
        
        # Validate paragraph index
        if paragraph_index >= len(doc.paragraphs):
            return f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})."
        
        paragraph = doc.paragraphs[paragraph_index]
        
        # Determine reference symbol
        if symbol is None:
            if note_type == "footnote":
                symbol = "¹"  # Unicode superscript 1 - could be enhanced with auto-numbering
            else:  # endnote
                symbol = "†"  # Unicode dagger symbol
        
        # Add note reference to paragraph
        if position == "beginning" and paragraph.runs:
            # Insert at beginning by modifying first run
            first_run = paragraph.runs[0]
            first_run.text = symbol + first_run.text
            first_run.font.superscript = True
        else:
            # Add at end (default behavior)
            note_run = paragraph.add_run(symbol)
            note_run.font.superscript = True
        
        # Handle note section creation/updating
        if note_type == "footnote":
            # Find or create footnotes section
            footnote_section_found = False
            for p in doc.paragraphs:
                if p.text.startswith("Footnotes:"):
                    footnote_section_found = True
                    break
            
            if not footnote_section_found:
                # Add footnotes section
                doc.add_paragraph("\\n").add_run()
                footnotes_heading = doc.add_paragraph("Footnotes:")
                footnotes_heading.bold = True
            
            # Add footnote text
            footnote_para = doc.add_paragraph(f"{symbol} {note_text}")
            if "Footnote Text" in doc.styles:
                footnote_para.style = "Footnote Text"
            else:
                footnote_para.style = "Normal"
        
        else:  # endnote
            # Find or create endnotes section
            endnotes_heading_found = False
            for para in doc.paragraphs:
                if para.text in ["Endnotes:", "ENDNOTES"]:
                    endnotes_heading_found = True
                    break
            
            if not endnotes_heading_found:
                # Add page break before endnotes section
                doc.add_page_break()
                doc.add_heading("Endnotes:", level=1)
            
            # Add endnote text
            endnote_para = doc.add_paragraph(f"{symbol} {note_text}")
            if "Endnote Text" in doc.styles:
                endnote_para.style = "Endnote Text"
            else:
                endnote_para.style = "Normal"
        
        doc.save(filename)
        
        return f"{note_type.capitalize()} added to paragraph {paragraph_index} in {filename}"
    
    except Exception as e:
        return f"Failed to add {note_type}: {str(e)}"

