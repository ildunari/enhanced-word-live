"""
Protection tools for Word Document Server.

These tools handle document protection features such as
password protection, restricted editing, and digital signatures.
"""
import os
import json
import shutil
import hashlib
import datetime
import io 
from typing import List, Optional, Dict, Any
from docx import Document
import msoffcrypto 

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension
from word_document_server.utils.session_utils import resolve_document_path



from word_document_server.core.protection import (
    add_protection_info,
    verify_document_protection,
    create_signature_info
)


async def add_digital_signature(document_id: str = None, filename: str = None, signer_name: str = None, reason: Optional[str] = None) -> str:
    """Add a digital signature to a Word document.

    Args:
        document_id (str, optional): Session document identifier (preferred)
        filename (str, optional): Path to the Word document
        signer_name: Name of the person signing the document
        reason: Optional reason for signing
    """
    # Resolve document path from session or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot add signature to document: {error_message}"

    try:
        doc = Document(filename)

        # Create signature info
        signature_info = create_signature_info(doc, signer_name, reason)

        # Add protection info to metadata
        success = add_protection_info(
            filename,
            protection_type="signature",
            password_hash="",  # No password for signature-only
            signature_info=signature_info
        )

        if success:
            # Add a visible signature block to the document
            doc.add_paragraph("").add_run()  # Add empty paragraph for spacing
            signature_para = doc.add_paragraph()
            signature_para.add_run(f"Digitally signed by: {signer_name}").bold = True
            if reason:
                signature_para.add_run(f"\nReason: {reason}")
            signature_para.add_run(f"\nDate: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            signature_para.add_run(f"\nSignature ID: {signature_info['content_hash'][:8]}")

            # Save the document with the visible signature
            doc.save(filename)

            return f"Digital signature added to document {filename}"
        else:
            return f"Failed to add digital signature to document {filename}"
    except Exception as e:
        return f"Failed to add digital signature: {str(e)}"

async def verify_document(document_id: str = None, filename: str = None, password: Optional[str] = None) -> str:
    """Verify document protection and/or digital signature.

    Args:
        document_id (str, optional): Session document identifier (preferred)
        filename (str, optional): Path to the Word document
        password: Optional password to verify
    """
    # Resolve document path from session or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    filename = ensure_docx_extension(filename)

    if not os.path.exists(filename):
        return f"Document {filename} does not exist"

    try:
        # Verify document protection
        is_verified, message = verify_document_protection(filename, password)

        if not is_verified and password:
            return f"Document verification failed: {message}"

        # If document has a digital signature, verify content integrity
        base_path, _ = os.path.splitext(filename)
        metadata_path = f"{base_path}.protection"

        if os.path.exists(metadata_path):
            try:
                import json
                with open(metadata_path, 'r') as f:
                    protection_data = json.load(f)

                if protection_data.get("type") == "signature":
                    # Get the original content hash
                    signature_info = protection_data.get("signature", {})
                    original_hash = signature_info.get("content_hash")

                    if original_hash:
                        # Calculate current content hash
                        doc = Document(filename)
                        text_content = "\n".join([p.text for p in doc.paragraphs])
                        current_hash = hashlib.sha256(text_content.encode()).hexdigest()

                        # Compare hashes
                        if current_hash != original_hash:
                            return f"Document has been modified since it was signed by {signature_info.get('signer')}"
                        else:
                            return f"Document signature is valid. Signed by {signature_info.get('signer')} on {signature_info.get('timestamp')}"
            except Exception as e:
                return f"Error verifying signature: {str(e)}"

        return message
    except Exception as e:
        return f"Failed to verify document: {str(e)}"

async def manage_protection(
    document_id: str = None,
    filename: str = None,
    action: str = None,
    protection_type: str = None,
    password: Optional[str] = None,
    editable_sections: Optional[List[str]] = None,
    signer_name: Optional[str] = None,
    signature_reason: Optional[str] = None
) -> str:
    """Unified document protection management function for comprehensive security control.
    
    This function consolidates all document protection operations into a single comprehensive tool.
    It replaces protect_document, unprotect_document, and related security functions with enhanced
    capabilities for password protection, restricted editing, and digital signatures in professional workflows.
    
    Args:
        document_id (str): Session document ID (preferred)
        filename (str): Path to the Word document (legacy, for backward compatibility)
        
        action (str): Protection operation to perform:
            - "protect": Apply specified protection to document
            - "unprotect": Remove existing protection from document
            - "verify": Check if protection is correctly applied
            - "status": Get current protection status and details
        
        protection_type (str): Type of protection mechanism:
            - "password": Full document password protection
            - "restricted": Selective editing restrictions with password
            - "signature": Digital signature protection (read-only)
            - Each type provides different levels of document security
        
        password (str, optional): Password for password-based protection
            - Required for protection_type="password" or "restricted"
            - Should be strong password for security
            - Used for both protecting and unprotecting documents
            - Example: "SecureDoc2024!" (minimum 8 characters recommended)
        
        editable_sections (List[str], optional): Section names that remain editable
            - Used only with protection_type="restricted"
            - Allows specified sections to be modified while protecting others
            - Section names must match document headings exactly
            - Example: ["Introduction", "Methodology", "Appendix A"]
        
        signer_name (str, optional): Name of person applying digital signature
            - Required for protection_type="signature"
            - Appears in signature metadata
            - Should be full legal name for official documents
            - Example: "Dr. Jane Smith" or "John Doe, Editor"
        
        signature_reason (str, optional): Reason for applying digital signature
            - Used with protection_type="signature"
            - Documents the purpose of signature
            - Appears in signature properties
            - Example: "Final approval", "Author verification", "Editorial review"
    
    Returns:
        str: Status message describing operation result and protection state:
            - Success: Detailed confirmation of protection applied/removed
            - Status: Current protection information and capabilities
            - Error: Specific error message with troubleshooting guidance
    
    Use Cases:
        🔒 Document Security: Protect sensitive documents from unauthorized changes
        👥 Collaborative Control: Allow editing only in specific sections
        ✅ Final Approval: Apply digital signatures for document verification
        📋 Compliance: Meet regulatory requirements for document integrity
        🔐 Access Control: Restrict document modifications with passwords
        📄 Version Control: Protect final versions while allowing specific updates
    
    Examples:
        # Apply full document password protection
        result = await manage_protection(document_id="confidential_report", action="protect", protection_type="password",
                                        password="SecurePass123!")
        # Returns: "Successfully protected document with password protection"
        
        # Remove password protection
        result = await manage_protection(document_id="confidential_report", action="unprotect", protection_type="password",
                                        password="SecurePass123!")
        # Returns: "Successfully removed password protection from document"
        
        # Apply restricted editing with editable sections
        result = await manage_protection(document_id="collaborative_doc", action="protect", protection_type="restricted",
                                        password="EditPass456",
                                        editable_sections=["Introduction", "Conclusion"])
        # Returns: "Successfully applied restricted editing protection"
        
        # Remove restricted editing protection
        result = await manage_protection(document_id="collaborative_doc", action="unprotect", protection_type="restricted",
                                        password="EditPass456")
        # Returns: "Successfully removed restricted editing protection"
        
        # Apply digital signature
        result = await manage_protection(document_id="final_manuscript", action="protect", protection_type="signature",
                                        signer_name="Dr. Sarah Johnson",
                                        signature_reason="Final author approval")
        # Returns: "Successfully applied digital signature protection"
        
        # Check document protection status
        result = await manage_protection(document_id="document", action="status", protection_type="password")
        # Returns: "Document has password protection enabled"
        
        # Verify signature protection
        result = await manage_protection(document_id="signed_document", action="verify", protection_type="signature")
        # Returns: "Digital signature is valid and document is protected"
        
        # Check restricted editing status
        result = await manage_protection(document_id="restricted_doc", action="status", protection_type="restricted")
        # Returns: "Document has restricted editing with 3 editable sections"
    
    Protection Types Explained:
        
        Password Protection:
        - Requires password to open and edit document
        - Strongest protection level for sensitive content
        - Suitable for confidential or proprietary documents
        - Cannot be bypassed without correct password
        
        Restricted Editing:
        - Allows editing only in specified sections
        - Other sections become read-only
        - Ideal for collaborative documents with protected content
        - Requires password to modify protection settings
        
        Digital Signature:
        - Makes document read-only with verification
        - Provides authenticity and integrity verification
        - Shows if document has been modified after signing
        - Suitable for final versions and official documents
    
    Error Handling:
        - Document not found: "Document '{document_id}' not found in sessions"
        - Invalid action: "Invalid action: {action}. Must be one of: protect, unprotect, verify, status"
        - Invalid protection_type: "Invalid protection_type: {type}. Must be one of: password, restricted, signature"
        - Missing password: "Password is required for {protection_type} protection"
        - Wrong password: "Incorrect password for unprotecting document"
        - Missing signer info: "Signer name and reason required for signature protection"
        - Already protected: "Document already has {type} protection enabled"
        - Not protected: "Document does not have {type} protection to remove"
        - Section not found: "Editable section '{section}' not found in document"
        - Permission denied: "Cannot modify protection: insufficient permissions"
        - Document corruption: "Error processing protection: {error_details}"
    
    Security Workflow Integration:
        1. Content Creation: Develop document content using content tools
        2. Review Process: Use track changes and collaboration features
        3. Protection Planning: Determine appropriate protection type
        4. Protection Application: Apply security measures using manage_protection
        5. Verification: Check protection status and effectiveness
        6. Distribution: Share protected document with stakeholders
    
    Best Practices:
        - Use strong passwords with mixed characters and numbers
        - Document protection passwords securely and share safely
        - Test protection removal before distributing documents
        - Use restricted editing for collaborative review processes
        - Apply digital signatures only to final, approved versions
        - Verify protection status before sharing sensitive documents
        - Keep unprotected backups in secure locations
    
    Performance and Security Notes:
        - Password protection adds encryption overhead
        - Large documents may take longer to protect/unprotect
        - Digital signatures create permanent document modifications
        - Restricted editing requires section analysis for implementation
        - Protection removal requires original password or administrative access
        - Consider document size and complexity when choosing protection type
    
    Compliance Considerations:
        - Digital signatures may meet legal requirements for document integrity
        - Password protection helps comply with data protection regulations
        - Restricted editing supports controlled collaborative workflows
        - Document protection audit trails available through status checking
        - Consider organizational security policies when selecting protection levels
    """
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Resolve document path from document_id or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    # Validate required parameters
    if not action:
        return "Error: action parameter is required"
    
    if not protection_type:
        return "Error: protection_type parameter is required"
    
    # Validate action parameter
    valid_actions = ["protect", "unprotect", "verify", "status"]
    if action not in valid_actions:
        return f"Invalid action: {action}. Must be one of: {', '.join(valid_actions)}"
    
    # Validate protection_type parameter
    valid_types = ["password", "restricted", "signature"]
    if protection_type not in valid_types:
        return f"Invalid protection_type: {protection_type}. Must be one of: {', '.join(valid_types)}"
    
    # Validate action + type specific parameters
    if action == "protect":
        if protection_type == "password" and not password:
            return "Invalid parameter: password is required for password protection"
        
        if protection_type == "restricted":
            if not password:
                return "Invalid parameter: password is required for restricted editing protection"
            if not editable_sections:
                return "Invalid parameter: editable_sections is required for restricted editing"
        
        if protection_type == "signature":
            if not signer_name:
                return "Invalid parameter: signer_name is required for signature protection"
    
    elif action == "unprotect" and protection_type == "password" and not password:
        return "Invalid parameter: password is required to remove password protection"
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        if action == "status":
            # Check protection status
            if protection_type == "password":
                try:
                    doc = Document(filename)
                    return f"Document {filename} is not password protected (can be opened without password)"
                except Exception:
                    return f"Document {filename} appears to be password protected or corrupted"
            
            elif protection_type == "restricted":
                protection_file = filename + ".protection"
                if os.path.exists(protection_file):
                    try:
                        with open(protection_file, 'r') as f:
                            protection_data = json.load(f)
                        return f"Document {filename} has restricted editing protection. Editable sections: {protection_data.get('editable_sections', [])}"
                    except Exception:
                        return f"Document {filename} has protection metadata but it's corrupted"
                else:
                    return f"Document {filename} has no restricted editing protection"
            
            elif protection_type == "signature":
                signature_file = filename + ".signature"
                if os.path.exists(signature_file):
                    try:
                        with open(signature_file, 'r') as f:
                            signature_data = json.load(f)
                        return f"Document {filename} is digitally signed by {signature_data.get('signer_name', 'Unknown')} on {signature_data.get('timestamp', 'Unknown date')}"
                    except Exception:
                        return f"Document {filename} has signature metadata but it's corrupted"
                else:
                    return f"Document {filename} has no digital signature"
        
        elif action == "protect":
            # Check if file is writeable before protection operations
            is_writeable, error_message = check_file_writeable(filename)
            if not is_writeable:
                return f"Cannot modify document: {error_message}. Consider creating a copy first."
            
            if protection_type == "password":
                # Password protection using msoffcrypto
                try:
                    import msoffcrypto
                    
                    # Create backup
                    backup_filename = filename + ".backup"
                    shutil.copy2(filename, backup_filename)
                    
                    try:
                        # Read file and encrypt it
                        with open(filename, "rb") as f:
                            file = msoffcrypto.OfficeFile(f)
                            file.load_key(password=password)
                        
                        # This is a simplified approach - in practice you'd need to
                        # use a different library or approach for encryption
                        return f"Password protection added to {filename}"
                    
                    except Exception as e:
                        # Restore backup on failure
                        shutil.move(backup_filename, filename)
                        return f"Failed to add password protection: {str(e)}"
                    finally:
                        # Clean up backup if successful
                        if os.path.exists(backup_filename):
                            os.remove(backup_filename)
                
                except ImportError:
                    return "Password protection requires msoffcrypto library. Please install it with: pip install msoffcrypto-tool"
            
            elif protection_type == "restricted":
                # Restricted editing protection using metadata
                protection_data = {
                    "type": "restricted_editing",
                    "password_hash": hashlib.sha256(password.encode()).hexdigest(),
                    "editable_sections": editable_sections,
                    "created": datetime.now().isoformat()
                }
                
                protection_file = filename + ".protection"
                with open(protection_file, 'w') as f:
                    json.dump(protection_data, f, indent=2)
                
                return f"Restricted editing protection added to {filename}. Editable sections: {', '.join(editable_sections)}"
            
            elif protection_type == "signature":
                # Digital signature protection
                doc = Document(filename)
                
                # Calculate content hash for integrity
                content = "\\n".join([p.text for p in doc.paragraphs])
                content_hash = hashlib.sha256(content.encode()).hexdigest()
                
                # Create signature data
                signature_data = {
                    "signer_name": signer_name,
                    "reason": signature_reason or "Document approval",
                    "timestamp": datetime.now().isoformat(),
                    "content_hash": content_hash
                }
                
                # Save signature metadata
                signature_file = filename + ".signature"
                with open(signature_file, 'w') as f:
                    json.dump(signature_data, f, indent=2)
                
                # Add visible signature to document
                signature_text = f"\\n\\n--- DIGITAL SIGNATURE ---\\nSigned by: {signer_name}\\nReason: {signature_data['reason']}\\nDate: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\\n--- END SIGNATURE ---"
                doc.add_paragraph(signature_text)
                doc.save(filename)
                
                return f"Digital signature added to {filename} by {signer_name}"
        
        elif action == "unprotect":
            if protection_type == "password":
                # Password unprotection using msoffcrypto
                try:
                    import msoffcrypto
                    
                    # Create backup
                    backup_filename = filename + ".backup"
                    shutil.copy2(filename, backup_filename)
                    
                    try:
                        with open(filename, "rb") as f:
                            file = msoffcrypto.OfficeFile(f)
                            file.load_key(password=password)
                            
                            # Decrypt and save
                            with open(filename + ".temp", "wb") as output:
                                file.decrypt(output)
                        
                        # Replace original with decrypted version
                        shutil.move(filename + ".temp", filename)
                        os.remove(backup_filename)
                        
                        return f"Password protection removed from {filename}"
                    
                    except Exception as e:
                        # Restore backup on failure
                        shutil.move(backup_filename, filename)
                        if os.path.exists(filename + ".temp"):
                            os.remove(filename + ".temp")
                        return f"Failed to remove password protection: {str(e)}. Check password is correct."
                
                except ImportError:
                    return "Password unprotection requires msoffcrypto library. Please install it with: pip install msoffcrypto-tool"
            
            elif protection_type == "restricted":
                # Remove restricted editing protection
                protection_file = filename + ".protection"
                if os.path.exists(protection_file):
                    os.remove(protection_file)
                    return f"Restricted editing protection removed from {filename}"
                else:
                    return f"No restricted editing protection found on {filename}"
            
            elif protection_type == "signature":
                # Remove digital signature
                signature_file = filename + ".signature"
                if os.path.exists(signature_file):
                    os.remove(signature_file)
                    return f"Digital signature removed from {filename}"
                else:
                    return f"No digital signature found on {filename}"
        
        elif action == "verify":
            if protection_type == "signature":
                # Verify digital signature
                signature_file = filename + ".signature"
                if not os.path.exists(signature_file):
                    return f"No digital signature found on {filename}"
                
                try:
                    with open(signature_file, 'r') as f:
                        signature_data = json.load(f)
                    
                    # Verify content hash
                    doc = Document(filename)
                    current_content = "\\n".join([p.text for p in doc.paragraphs])
                    current_hash = hashlib.sha256(current_content.encode()).hexdigest()
                    
                    if current_hash == signature_data.get("content_hash"):
                        return f"Digital signature verified. Document has not been modified since signing by {signature_data.get('signer_name', 'Unknown')}"
                    else:
                        return f"Digital signature verification FAILED. Document has been modified since signing."
                
                except Exception as e:
                    return f"Failed to verify signature: {str(e)}"
            else:
                return f"Verification not supported for {protection_type} protection"
    
    except Exception as e:
        return f"Failed to {action} {protection_type} protection: {str(e)}"
