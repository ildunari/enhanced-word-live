"""
Document creation and manipulation tools for Word Document Server.
"""
import os
import json
from typing import Dict, List, Optional, Any
from docx import Document

from word_document_server.utils.file_utils import check_file_writeable, ensure_docx_extension, create_document_copy
from word_document_server.utils.document_utils import get_document_properties, extract_document_text, get_document_structure
from word_document_server.utils.extended_document_utils import get_paragraph_text, find_text
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document with optional metadata.
    
    Args:
        filename: Name of the document to create (with or without .docx extension)
        title: Optional title for the document metadata
        author: Optional author for the document metadata
    """
    filename = ensure_docx_extension(filename)
    
    # Check if file is writeable
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot create document: {error_message}"
    
    try:
        doc = Document()
        
        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author
        
        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)
        
        # Save the document
        doc.save(filename)
        
        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}"


async def get_document_info(filename: str) -> str:
    """Get information about a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    if not os.path.exists(filename):
        return f"Document {filename} does not exist"
    
    try:
        properties = get_document_properties(filename)
        return json.dumps(properties, indent=2)
    except Exception as e:
        return f"Failed to get document info: {str(e)}"


async def get_text(
    document_id: str = None,
    filename: str = None,
    scope: str = "all",
    paragraph_index: Optional[int] = None,
    search_term: Optional[str] = None,
    start_paragraph: Optional[int] = None,
    end_paragraph: Optional[int] = None,
    include_formatting: bool = False,
    formatting_detail: str = "basic",
    max_results: int = 100,
    match_case: bool = True,
    whole_word: bool = False
) -> str:
    """Unified text extraction function combining document, paragraph, and search functionality.
    
    This function consolidates multiple text extraction operations into a single comprehensive tool.
    It replaces get_document_text, get_paragraph_text_from_document, and find_text_in_document.
    
    Args:
        document_id (str): Session document ID (preferred)
        filename (str): Path to the Word document (legacy, for backward compatibility)
        
        scope (str): Text extraction scope - determines what content is extracted:
            - "all": Extract entire document content (default)
            - "paragraph": Extract specific paragraph by index
            - "search": Search for specific text within document
            - "range": Extract paragraph range between start/end indices
        
        paragraph_index (int, optional): Zero-based paragraph index for scope="paragraph"
            - Example: 0 = first paragraph, 5 = sixth paragraph
            - Required when scope="paragraph"
        
        search_term (str, optional): Text pattern to search for when scope="search"
            - Supports partial matches unless whole_word=True
            - Case sensitivity controlled by match_case parameter
            - Required when scope="search"
        
        start_paragraph (int, optional): Starting paragraph index for scope="range"
            - Zero-based index, inclusive
            - If not provided with scope="range", starts from beginning
        
        end_paragraph (int, optional): Ending paragraph index for scope="range" 
            - Zero-based index, inclusive
            - If not provided with scope="range", continues to end
        
        include_formatting (bool): Whether to include formatting information in output
            - False: Plain text only (default)
            - True: Include font, style, and formatting details
        
        formatting_detail (str): Level of formatting information when include_formatting=True:
            - "basic": Font name, size, bold/italic status
            - "detailed": Basic + color, alignment, spacing
            - "comprehensive": Detailed + advanced formatting properties
        
        max_results (int): Maximum number of search results for scope="search" (default: 100)
            - Limits output size for large documents
            - Only affects search operations
        
        match_case (bool): Case-sensitive matching for scope="search" (default: True)
            - True: "Word" != "word"
            - False: "Word" == "word"
        
        whole_word (bool): Match whole words only for scope="search" (default: False)
            - True: "cat" won't match "catch"
            - False: "cat" will match "catch"
    
    Returns:
        str: Extracted content in format determined by scope and formatting options:
            - scope="all"|"paragraph"|"range": Plain text or JSON with formatting
            - scope="search": JSON object with matches, contexts, and formatting
            - Error message string if operation fails
    
    Use Cases:
        📄 Document Review: Extract full text for analysis or content review
        📝 Content Extraction: Get specific paragraphs for editing or citation
        🔍 Research: Search for key terms, findings, or references
        📊 Analysis: Extract formatted content for further processing
        📋 Academic Work: Find citations, extract methodology sections
        ✏️ Editing: Get paragraph ranges for revision or restructuring
    
    Examples:
        # Basic document extraction
        text = await get_text(document_id="main")
        # Returns: Full document text as string
        
        # Extract specific paragraph with formatting
        para = await get_text(document_id="thesis", scope="paragraph", paragraph_index=5, 
                             include_formatting=True, formatting_detail="detailed")
        # Returns: JSON with paragraph text and formatting details
        
        # Search for methodology section
        methods = await get_text(document_id="paper", scope="search", search_term="methodology",
                                include_formatting=True, max_results=5)
        # Returns: JSON with search matches and surrounding context
        
        # Extract conclusion section (paragraphs 45-50)
        conclusion = await get_text(document_id="report", scope="range", 
                                   start_paragraph=45, end_paragraph=50,
                                   include_formatting=True, formatting_detail="comprehensive")
        # Returns: JSON with paragraph range and comprehensive formatting
        
        # Case-insensitive search for citations
        citations = await get_text(document_id="document", scope="search", search_term="et al",
                                  match_case=False, whole_word=True, max_results=20)
        # Returns: JSON with all citation matches
        
        # Academic paper abstract extraction (assuming it's paragraph 2)
        abstract = await get_text(document_id="paper", scope="paragraph", paragraph_index=1)
        # Returns: Plain text of abstract paragraph
    
    Error Handling:
        - Document not found: Returns "Document '{document_id}' not found in sessions"
        - Invalid scope: Returns "Invalid scope: {scope}. Must be one of: all, paragraph, search, range"
        - Missing required parameters: Returns specific error for missing parameter
        - Invalid paragraph index: Returns "Paragraph index {index} is out of range"
        - Document corruption: Returns "Error reading document: {error_details}"
        - Permission issues: Returns "Cannot read document: {permission_error}"
    
    Performance Notes:
        - Large documents with include_formatting=True may take longer to process
        - Search operations are optimized but may be slower on very large documents
        - Formatting extraction adds processing time proportional to formatting_detail level
        - Consider using max_results to limit search output for performance
    """
    from word_document_server.utils.session_utils import resolve_document_path
    from word_document_server.session_manager import get_session_manager
    import io
    
    # Resolve document path from document_id or filename
    filename, error_msg = resolve_document_path(document_id, filename)
    if error_msg:
        return error_msg
    
    # Validate scope parameter
    valid_scopes = ["all", "paragraph", "search", "range"]
    if scope not in valid_scopes:
        return f"Invalid scope: {scope}. Must be one of: {', '.join(valid_scopes)}"
    
    # Validate formatting_detail parameter
    valid_details = ["basic", "detailed", "comprehensive"]
    if formatting_detail not in valid_details:
        return f"Invalid formatting_detail: {formatting_detail}. Must be one of: {', '.join(valid_details)}"
    
    # Validate scope-specific parameters
    if scope == "paragraph" and paragraph_index is None:
        return "Invalid parameter: paragraph_index is required when scope='paragraph'"
    
    if scope == "search" and not search_term:
        return "Invalid parameter: search_term is required when scope='search'"
    
    if scope == "range" and (start_paragraph is None or end_paragraph is None):
        return "Invalid parameter: both start_paragraph and end_paragraph are required when scope='range'"
    
    # Validate numeric parameters
    if paragraph_index is not None:
        try:
            paragraph_index = int(paragraph_index)
            if paragraph_index < 0:
                return "Invalid parameter: paragraph_index must be a non-negative integer"
        except (ValueError, TypeError):
            return "Invalid parameter: paragraph_index must be an integer"
    
    if start_paragraph is not None:
        try:
            start_paragraph = int(start_paragraph)
            if start_paragraph < 0:
                return "Invalid parameter: start_paragraph must be a non-negative integer"
        except (ValueError, TypeError):
            return "Invalid parameter: start_paragraph must be an integer"
    
    if end_paragraph is not None:
        try:
            end_paragraph = int(end_paragraph)
            if end_paragraph < 0:
                return "Invalid parameter: end_paragraph must be a non-negative integer"
        except (ValueError, TypeError):
            return "Invalid parameter: end_paragraph must be an integer"
    
    # Get session manager to check for live connections
    session_manager = get_session_manager()
    
    # --- LIVE EDITING LOGIC ---
    if document_id and session_manager.is_document_live(document_id):
        try:
            # Get current content from live document
            ooxml_response = await session_manager.send_live_request(document_id, "get_full_content")
            ooxml_content = ooxml_response.get("content")
            if not ooxml_content:
                return "Failed to retrieve content from live document"
            
            # Create Document object from live content
            doc = Document(io.BytesIO(ooxml_content.encode('utf-8')))
            print(f"[get_text] Operating in LIVE mode for document '{document_id}'")
            
        except Exception as e:
            return f"Failed to get live document content: {str(e)}"
    
    # --- FILE-BASED LOGIC (FALLBACK) ---
    else:
        if not os.path.exists(filename):
            return f"Document {filename} does not exist"
        
        # Load document from file
        doc = Document(filename)
        print(f"[get_text] Operating in FILE mode for {filename}")
    
    def extract_run_formatting(run, detail_level="basic"):
        """Extract formatting information from a run."""
        formatting = {
            "text": run.text,
        }
        
        if detail_level in ["basic", "detailed", "comprehensive"]:
            # Basic formatting
            formatting.update({
                "bold": run.bold,
                "italic": run.italic,
                "underline": run.underline,
            })
        
        if detail_level in ["detailed", "comprehensive"]:
            # Detailed formatting
            formatting.update({
                "font_name": run.font.name,
                "font_size": str(run.font.size) if run.font.size else None,
                "font_color": str(run.font.color.rgb) if run.font.color.rgb else None,
                "highlight_color": str(run.font.highlight_color) if run.font.highlight_color else None,
                "strike": run.font.strike,
                "double_strike": run.font.double_strike,
                "superscript": run.font.superscript,
                "subscript": run.font.subscript,
                "small_caps": run.font.small_caps,
                "all_caps": run.font.all_caps,
            })
        
        if detail_level == "comprehensive":
            # Comprehensive formatting
            formatting.update({
                "font_color_theme": str(run.font.color.theme_color) if run.font.color.theme_color else None,
                "font_color_brightness": run.font.color.brightness if hasattr(run.font.color, 'brightness') else None,
                "emboss": run.font.emboss,
                "imprint": run.font.imprint,
                "outline": run.font.outline,
                "shadow": run.font.shadow,
                "snap_to_grid": run.font.snap_to_grid,
                "spec_vanish": run.font.spec_vanish,
                "web_hidden": run.font.web_hidden,
                "cs_bold": run.font.cs_bold,
                "cs_italic": run.font.cs_italic,
                "east_asia_font": run.font.name_east_asia,
                "complex_script_font": run.font.name_cs,
            })
        
        # Clean up None values for cleaner output
        return {k: v for k, v in formatting.items() if v is not None}
    
    def extract_paragraph_formatting(paragraph, detail_level="basic"):
        """Extract formatting information from a paragraph."""
        formatting = {}
        
        if detail_level in ["basic", "detailed", "comprehensive"]:
            # Basic paragraph formatting
            formatting.update({
                "style": paragraph.style.name if paragraph.style else None,
                "alignment": str(paragraph.alignment) if paragraph.alignment else None,
            })
        
        if detail_level in ["detailed", "comprehensive"]:
            # Detailed paragraph formatting
            paragraph_format = paragraph.paragraph_format
            formatting.update({
                "left_indent": str(paragraph_format.left_indent) if paragraph_format.left_indent else None,
                "right_indent": str(paragraph_format.right_indent) if paragraph_format.right_indent else None,
                "first_line_indent": str(paragraph_format.first_line_indent) if paragraph_format.first_line_indent else None,
                "space_before": str(paragraph_format.space_before) if paragraph_format.space_before else None,
                "space_after": str(paragraph_format.space_after) if paragraph_format.space_after else None,
                "line_spacing": str(paragraph_format.line_spacing) if paragraph_format.line_spacing else None,
            })
        
        if detail_level == "comprehensive":
            # Comprehensive paragraph formatting
            paragraph_format = paragraph.paragraph_format
            formatting.update({
                "keep_together": paragraph_format.keep_together,
                "keep_with_next": paragraph_format.keep_with_next,
                "page_break_before": paragraph_format.page_break_before,
                "widow_control": paragraph_format.widow_control,
                "line_spacing_rule": str(paragraph_format.line_spacing_rule) if paragraph_format.line_spacing_rule else None,
                "tab_stops": [{"position": str(tab.position), "alignment": str(tab.alignment), "leader": str(tab.leader)} 
                             for tab in paragraph_format.tab_stops] if paragraph_format.tab_stops else []
            })
        
        # Clean up None values for cleaner output
        return {k: v for k, v in formatting.items() if v is not None}
    
    try:
        if scope == "all":
            # Original get_document_text functionality with optional formatting
            if not include_formatting:
                # For live documents, we need to extract text from the already loaded doc
                if document_id and session_manager.is_document_live(document_id):
                    return "\n".join([p.text for p in doc.paragraphs])
                else:
                    return extract_document_text(filename)
            else:
                # doc is already loaded above (either from live or file)
                result = {
                    "document_text": "",
                    "paragraphs": [],
                    "formatting_detail": formatting_detail
                }
                
                for i, paragraph in enumerate(doc.paragraphs):
                    para_info = {
                        "index": i,
                        "text": paragraph.text,
                        "paragraph_formatting": extract_paragraph_formatting(paragraph, formatting_detail),
                        "runs": []
                    }
                    
                    for run in paragraph.runs:
                        if run.text.strip():  # Only include runs with actual text
                            para_info["runs"].append(extract_run_formatting(run, formatting_detail))
                    
                    result["paragraphs"].append(para_info)
                    result["document_text"] += paragraph.text + "\n"
                
                return json.dumps(result, indent=2)
        
        elif scope == "paragraph":
            # Original get_paragraph_text_from_document functionality with enhanced formatting
            # doc is already loaded above (either from live or file)
            
            # Validate paragraph index
            if paragraph_index >= len(doc.paragraphs):
                return f"Invalid paragraph index: {paragraph_index}. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})"
            
            paragraph = doc.paragraphs[paragraph_index]
            
            if not include_formatting:
                result = get_paragraph_text(filename, paragraph_index)
                return json.dumps(result, indent=2)
            else:
                result = {
                    "paragraph_index": paragraph_index,
                    "text": paragraph.text,
                    "paragraph_formatting": extract_paragraph_formatting(paragraph, formatting_detail),
                    "runs": [],
                    "formatting_detail": formatting_detail
                }
                
                for run in paragraph.runs:
                    if run.text.strip():  # Only include runs with actual text
                        result["runs"].append(extract_run_formatting(run, formatting_detail))
                
                return json.dumps(result, indent=2)
        
        elif scope == "search":
            # Original find_text_in_document functionality with enhanced formatting
            if not include_formatting:
                # For live documents, we can't use find_text(filename, ...) so we'll process the doc directly
                if document_id and session_manager.is_document_live(document_id):
                    # Simple text search for live documents when formatting is not needed
                    occurrences = []
                    search_lower = search_term.lower() if not match_case else search_term
                    
                    for para_idx, paragraph in enumerate(doc.paragraphs):
                        para_text = paragraph.text
                        search_text = para_text.lower() if not match_case else para_text
                        
                        if whole_word:
                            import re
                            pattern = r'' + re.escape(search_lower) + r''
                            matches = list(re.finditer(pattern, search_text, re.IGNORECASE if not match_case else 0))
                        else:
                            matches = []
                            start = 0
                            while True:
                                pos = search_text.find(search_lower, start)
                                if pos == -1:
                                    break
                                matches.append(type('Match', (), {'start': lambda: pos, 'end': lambda: pos + len(search_term)})())
                                start = pos + 1
                        
                        for match in matches:
                            if len(occurrences) >= max_results:
                                break
                            
                            pos = match.start()
                            end_pos = match.end()
                            context_start = max(0, pos - 50)
                            context_end = min(len(para_text), end_pos + 50)
                            
                            occurrences.append({
                                "paragraph_index": para_idx,
                                "character_position": pos,
                                "matched_text": para_text[pos:end_pos],
                                "context": para_text[context_start:context_end]
                            })
                        
                        if len(occurrences) >= max_results:
                            break
                    
                    result = {
                        "query": search_term,
                        "total_count": len(occurrences),
                        "occurrences": occurrences,
                        "source": "live_document"
                    }
                else:
                    result = find_text(filename, search_term, match_case, whole_word)
                # Limit results if max_results is specified
                if "occurrences" in result and len(result["occurrences"]) > max_results:
                    result["occurrences"] = result["occurrences"][:max_results]
                    result["total_count"] = len(result["occurrences"])
                    result["truncated"] = True
                return json.dumps(result, indent=2)
            else:
                # doc is already loaded above (either from live or file)
                occurrences = []
                
                search_lower = search_term.lower() if not match_case else search_term
                
                for para_idx, paragraph in enumerate(doc.paragraphs):
                    para_text = paragraph.text
                    search_text = para_text.lower() if not match_case else para_text
                    
                    # Find all occurrences in this paragraph
                    start = 0
                    while True:
                        if whole_word:
                            # Simple whole word matching
                            import re
                            pattern = r'\b' + re.escape(search_lower) + r'\b'
                            match = re.search(pattern, search_text[start:], re.IGNORECASE if not match_case else 0)
                            if match:
                                pos = start + match.start()
                                end_pos = start + match.end()
                            else:
                                break
                        else:
                            pos = search_text.find(search_lower, start)
                            if pos == -1:
                                break
                            end_pos = pos + len(search_term)
                        
                        # Extract context and formatting
                        context_start = max(0, pos - 50)
                        context_end = min(len(para_text), end_pos + 50)
                        context = para_text[context_start:context_end]
                        
                        # Find which run contains this text and extract its formatting
                        char_count = 0
                        containing_run = None
                        run_formatting = {}
                        
                        for run in paragraph.runs:
                            if char_count <= pos < char_count + len(run.text):
                                containing_run = run
                                run_formatting = extract_run_formatting(run, formatting_detail)
                                break
                            char_count += len(run.text)
                        
                        occurrence = {
                            "paragraph_index": para_idx,
                            "character_position": pos,
                            "matched_text": para_text[pos:end_pos],
                            "context": context,
                            "paragraph_formatting": extract_paragraph_formatting(paragraph, formatting_detail),
                            "run_formatting": run_formatting
                        }
                        
                        occurrences.append(occurrence)
                        
                        if len(occurrences) >= max_results:
                            break
                        
                        start = end_pos
                    
                    if len(occurrences) >= max_results:
                        break
                
                result = {
                    "query": search_term,
                    "match_case": match_case,
                    "whole_word": whole_word,
                    "formatting_detail": formatting_detail,
                    "total_count": len(occurrences),
                    "truncated": len(occurrences) >= max_results,
                    "occurrences": occurrences
                }
                
                return json.dumps(result, indent=2)
        
        elif scope == "range":
            # New functionality: extract paragraph range with optional formatting
            # doc is already loaded above (either from live or file)
            
            # Validate range parameters
            if start_paragraph >= len(doc.paragraphs):
                return f"Invalid start_paragraph: {start_paragraph}. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})"
            
            if end_paragraph >= len(doc.paragraphs):
                return f"Invalid end_paragraph: {end_paragraph}. Document has {len(doc.paragraphs)} paragraphs (0-{len(doc.paragraphs)-1})"
            
            if start_paragraph > end_paragraph:
                return f"Invalid range: start_paragraph ({start_paragraph}) must be <= end_paragraph ({end_paragraph})"
            
            # Extract text from paragraph range
            paragraphs = doc.paragraphs[start_paragraph:end_paragraph + 1]
            
            if not include_formatting:
                text_parts = []
                for i, paragraph in enumerate(paragraphs):
                    actual_index = start_paragraph + i
                    text_parts.append(f"[Paragraph {actual_index}] {paragraph.text}")
                
                return "\n".join(text_parts)
            else:
                result = {
                    "start_paragraph": start_paragraph,
                    "end_paragraph": end_paragraph,
                    "formatting_detail": formatting_detail,
                    "paragraphs": []
                }
                
                for i, paragraph in enumerate(paragraphs):
                    actual_index = start_paragraph + i
                    para_info = {
                        "index": actual_index,
                        "text": paragraph.text,
                        "paragraph_formatting": extract_paragraph_formatting(paragraph, formatting_detail),
                        "runs": []
                    }
                    
                    for run in paragraph.runs:
                        if run.text.strip():  # Only include runs with actual text
                            para_info["runs"].append(extract_run_formatting(run, formatting_detail))
                    
                    result["paragraphs"].append(para_info)
                
                return json.dumps(result, indent=2)
        
    except Exception as e:
        return f"Failed to extract text: {str(e)}"



async def get_document_outline(filename: str) -> str:
    """Get the structure of a Word document.
    
    Args:
        filename: Path to the Word document
    """
    filename = ensure_docx_extension(filename)
    
    structure = get_document_structure(filename)
    return json.dumps(structure, indent=2)


async def list_available_documents(directory: str = ".") -> str:
    """List all .docx files in the specified directory.
    
    Args:
        directory: Directory to search for Word documents
    """
    try:
        if not os.path.exists(directory):
            return f"Directory {directory} does not exist"
        
        docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
        
        if not docx_files:
            return f"No Word documents found in {directory}"
        
        result = f"Found {len(docx_files)} Word documents in {directory}:\n"
        for file in docx_files:
            file_path = os.path.join(directory, file)
            size = os.path.getsize(file_path) / 1024  # KB
            result += f"- {file} ({size:.2f} KB)\n"
        
        return result
    except Exception as e:
        return f"Failed to list documents: {str(e)}"


async def copy_document(source_filename: str, destination_filename: Optional[str] = None) -> str:
    """Create a copy of a Word document.
    
    Args:
        source_filename: Path to the source document
        destination_filename: Optional path for the copy. If not provided, a default name will be generated.
    """
    source_filename = ensure_docx_extension(source_filename)
    
    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)
    
    success, message, new_path = create_document_copy(source_filename, destination_filename)
    if success:
        return message
    else:
        return f"Failed to copy document: {message}"


async def merge_documents(target_filename: str, source_filenames: List[str], add_page_breaks: bool = True) -> str:
    """Merge multiple Word documents into a single document.
    
    Args:
        target_filename: Path to the target document (will be created or overwritten)
        source_filenames: List of paths to source documents to merge
        add_page_breaks: If True, add page breaks between documents
    """
    from word_document_server.core.tables import copy_table
    
    target_filename = ensure_docx_extension(target_filename)
    
    # Check if target file is writeable
    is_writeable, error_message = check_file_writeable(target_filename)
    if not is_writeable:
        return f"Cannot create target document: {error_message}"
    
    # Validate all source documents exist
    missing_files = []
    for filename in source_filenames:
        doc_filename = ensure_docx_extension(filename)
        if not os.path.exists(doc_filename):
            missing_files.append(doc_filename)
    
    if missing_files:
        return f"Cannot merge documents. The following source files do not exist: {', '.join(missing_files)}"
    
    try:
        # Create a new document for the merged result
        target_doc = Document()
        
        # Process each source document
        for i, filename in enumerate(source_filenames):
            doc_filename = ensure_docx_extension(filename)
            source_doc = Document(doc_filename)
            
            # Add page break between documents (except before the first one)
            if add_page_breaks and i > 0:
                target_doc.add_page_break()
            
            # Copy all paragraphs
            for paragraph in source_doc.paragraphs:
                # Create a new paragraph with the same text and style
                new_paragraph = target_doc.add_paragraph(paragraph.text)
                new_paragraph.style = target_doc.styles['Normal']  # Default style
                
                # Try to match the style if possible
                try:
                    if paragraph.style and paragraph.style.name in target_doc.styles:
                        new_paragraph.style = target_doc.styles[paragraph.style.name]
                except:
                    pass
                
                # Copy run formatting
                for i, run in enumerate(paragraph.runs):
                    if i < len(new_paragraph.runs):
                        new_run = new_paragraph.runs[i]
                        # Copy basic formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Font size if specified
                        if run.font.size:
                            new_run.font.size = run.font.size
            
            # Copy all tables
            for table in source_doc.tables:
                copy_table(table, target_doc)
        
        # Save the merged document
        target_doc.save(target_filename)
        return f"Successfully merged {len(source_filenames)} documents into {target_filename}"
    except Exception as e:
        return f"Failed to merge documents: {str(e)}"


def document_utility(
    action: str,
    document_id: str = None,
    filename: str = None,
    directory: str = "."
) -> str:
    """Unified document utility function for document information operations.
    
    This consolidated tool replaces 3 individual document information functions with a single
    action-based interface, reducing tool count while preserving 100% functionality.
    
    Args:
        action (str): Document utility operation to perform:
            - "info": Get document metadata and properties (requires document_id or filename)
            - "outline": Get document structure and outline (requires document_id or filename)
            - "list_files": List available Word documents in directory
        document_id (str, optional): Session document identifier (preferred for info/outline)
        filename (str, optional): Path to Word document (required for "info" and "outline" if no document_id)
        directory (str, optional): Directory to search (for "list_files" action, defaults to ".")
        
    Returns:
        str: Operation result as formatted string or JSON
        
    Examples:
        # Get document information (session-based)
        document_utility("info", document_id="main")
        
        # Get document structure/outline (legacy filename)
        document_utility("outline", filename="research_paper.docx")
        
        # List Word documents in current directory
        document_utility("list_files")
        
        # List Word documents in specific directory
        document_utility("list_files", "", "/Users/john/Documents")
    """
    import asyncio
    from word_document_server.utils.session_utils import resolve_document_path
    
    # Validate action parameter
    valid_actions = ["info", "outline", "list_files"]
    if action not in valid_actions:
        return f"Invalid action: {action}. Must be one of: {', '.join(valid_actions)}"
    
    # Resolve document path for info/outline actions
    if action in ["info", "outline"]:
        filename, error_msg = resolve_document_path(document_id, filename)
        if error_msg:
            return error_msg
    
    # Delegate to appropriate original function based on action
    try:
        import asyncio
        
        if action == "info":
            # Check if we're in an event loop
            try:
                loop = asyncio.get_running_loop()
                # We're in a running loop, create a task
                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(asyncio.run, get_document_info(filename))
                    return future.result()
            except RuntimeError:
                # No running loop, safe to use asyncio.run
                return asyncio.run(get_document_info(filename))
            
        elif action == "outline":
            try:
                loop = asyncio.get_running_loop()
                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(asyncio.run, get_document_outline(filename))
                    return future.result()
            except RuntimeError:
                return asyncio.run(get_document_outline(filename))
            
        elif action == "list_files":
            search_dir = directory if directory else "."
            try:
                loop = asyncio.get_running_loop()
                import concurrent.futures
                with concurrent.futures.ThreadPoolExecutor() as executor:
                    future = executor.submit(asyncio.run, list_available_documents(search_dir))
                    return future.result()
            except RuntimeError:
                return asyncio.run(list_available_documents(search_dir))
            
    except Exception as e:
        return f"Error in document_utility: {str(e)}"
