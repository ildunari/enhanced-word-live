#!/usr/bin/env python3
"""
Test script for the Live Editing Architecture upgrade.

This script tests the basic functionality of the upgraded Enhanced Word MCP Server
with live editing capabilities.
"""

import os
import sys
import asyncio
import json
import tempfile
from pathlib import Path

# Add the project root to the Python path
project_root = Path(__file__).parent
sys.path.insert(0, str(project_root))

from word_document_server.session_manager import get_session_manager
from word_document_server.tools.document_tools import create_document, get_text
from word_document_server.tools.content_tools import enhanced_search_and_replace
from docx import Document

def test_session_manager_live_capabilities():
    """Test that the session manager has live editing capabilities."""
    print("🧪 Testing Session Manager Live Capabilities...")
    
    session_manager = get_session_manager()
    
    # Check that new methods exist
    assert hasattr(session_manager, 'register_live_connection'), "register_live_connection method missing"
    assert hasattr(session_manager, 'unregister_live_connection'), "unregister_live_connection method missing"
    assert hasattr(session_manager, 'is_document_live'), "is_document_live method missing"
    assert hasattr(session_manager, 'send_live_request'), "send_live_request method missing"
    assert hasattr(session_manager, 'handle_live_response'), "handle_live_response method missing"
    
    print("✅ Session Manager has all required live editing methods")
    return True

def test_document_handle_extensions():
    """Test that DocumentHandle has been extended with WebSocket support."""
    print("🧪 Testing DocumentHandle Extensions...")
    
    # Create a temporary document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        tmp_path = tmp_file.name
    
    try:
        # Create a simple document
        doc = Document()
        doc.add_paragraph("Test document for live editing")
        doc.save(tmp_path)
        
        # Test document session creation
        session_manager = get_session_manager()
        result = session_manager.open_document("test_doc", tmp_path)
        
        if "Successfully opened" not in result:
            print(f"❌ Failed to open document: {result}")
            return False
        
        # Get the document handle and test live capabilities
        handle = session_manager.get_document("test_doc")
        assert handle is not None, "Document handle is None"
        assert hasattr(handle, 'websocket_connection'), "websocket_connection attribute missing"
        assert hasattr(handle, 'is_live'), "is_live property missing"
        assert hasattr(handle, 'register_websocket'), "register_websocket method missing"
        assert hasattr(handle, 'unregister_websocket'), "unregister_websocket method missing"
        
        # Test is_live property
        assert handle.is_live == False, "Document should not be live initially"
        
        # Clean up
        session_manager.close_document("test_doc")
        
        print("✅ DocumentHandle has all required live editing capabilities")
        return True
        
    finally:
        # Clean up temporary file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

async def test_enhanced_tools_async_support():
    """Test that enhanced tools support async operations."""
    print("🧪 Testing Enhanced Tools Async Support...")
    
    # Create a temporary document
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        tmp_path = tmp_file.name
    
    try:
        # Create a test document
        doc = Document()
        doc.add_paragraph("This is a test document for live editing capabilities.")
        doc.add_paragraph("It contains multiple paragraphs with various content.")
        doc.add_paragraph("We will test search and replace functionality.")
        doc.save(tmp_path)
        
        # Test get_text async function
        result = await get_text(filename=tmp_path, scope="all")
        assert "test document" in result.lower(), "get_text did not return expected content"
        
        # Test enhanced_search_and_replace async function  
        result = await enhanced_search_and_replace(
            filename=tmp_path,
            find_text="test document",
            replace_text="TEST DOCUMENT",
            match_case=False
        )
        assert "File mode:" in result, "enhanced_search_and_replace should indicate file mode"
        assert "Replaced" in result, "enhanced_search_and_replace should report replacement"
        
        print("✅ Enhanced tools support async operations")
        return True
        
    except Exception as e:
        print(f"❌ Async tools test failed: {e}")
        return False
        
    finally:
        # Clean up temporary file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

def test_websocket_imports():
    """Test that WebSocket-related imports work."""
    print("🧪 Testing WebSocket Imports...")
    
    try:
        import websockets
        import asyncio
        import json
        import threading
        
        # Test that main.py imports work
        from word_document_server.main import websocket_handler, run_websocket_server
        
        print("✅ All WebSocket imports successful")
        return True
        
    except ImportError as e:
        print(f"❌ WebSocket import failed: {e}")
        return False

def test_word_addin_structure():
    """Test that Word Add-in project structure was created."""
    print("🧪 Testing Word Add-in Project Structure...")
    
    addin_path = project_root / "word-live-addin"
    
    required_files = [
        "manifest.xml",
        "src/taskpane/taskpane.html",
        "src/taskpane/taskpane.css", 
        "src/taskpane/taskpane.js",
        "package.json",
        "README.md"
    ]
    
    for file_path in required_files:
        full_path = addin_path / file_path
        if not full_path.exists():
            print(f"❌ Missing Add-in file: {file_path}")
            return False
    
    # Check that manifest has the correct GUID format
    manifest_path = addin_path / "manifest.xml"
    with open(manifest_path, 'r') as f:
        manifest_content = f.read()
        if "<Id>" not in manifest_content or "</Id>" not in manifest_content:
            print("❌ Manifest missing GUID")
            return False
    
    print("✅ Word Add-in project structure is complete")
    return True

async def run_all_tests():
    """Run all tests and report results."""
    print("🚀 Starting Live Editing Architecture Upgrade Tests\n")
    
    tests = [
        ("Session Manager Live Capabilities", test_session_manager_live_capabilities),
        ("DocumentHandle Extensions", test_document_handle_extensions),
        ("Enhanced Tools Async Support", test_enhanced_tools_async_support),
        ("WebSocket Imports", test_websocket_imports),
        ("Word Add-in Structure", test_word_addin_structure),
    ]
    
    passed = 0
    total = len(tests)
    
    for test_name, test_func in tests:
        try:
            if asyncio.iscoroutinefunction(test_func):
                result = await test_func()
            else:
                result = test_func()
            
            if result:
                passed += 1
            print()  # Add spacing between tests
            
        except Exception as e:
            print(f"❌ {test_name} failed with exception: {e}\n")
    
    print("=" * 60)
    print(f"📊 Test Results: {passed}/{total} tests passed")
    
    if passed == total:
        print("🎉 All tests passed! Live editing upgrade is successful.")
        print("\n📋 Next Steps:")
        print("1. Install Word Add-in dependencies: cd word-live-addin && npm install")
        print("2. Start the MCP server: python -m word_document_server.main")
        print("3. Start the Word Add-in: cd word-live-addin && npm start")
        print("4. Load the Add-in in Microsoft Word and test live editing")
        return True
    else:
        print("❌ Some tests failed. Please review the errors above.")
        return False

if __name__ == "__main__":
    # Run the test suite
    success = asyncio.run(run_all_tests())
    sys.exit(0 if success else 1)